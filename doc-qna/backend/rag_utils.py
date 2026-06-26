# import json
# import logging
# import os
# import re
# import time
# from typing import Any, Dict, Iterable, List, Optional, Tuple

# from store_runtime_config import merge_store_runtime_config

# # Use the same logger family as the rest of the app so retrieval-chunk logs
# # end up in `data/logs/YYYY-MM-DD.log` (configured in qna_service.py).
# logger = logging.getLogger("yuktra_qna.rag")


# def pipeline_log_preview(text: Optional[str], *, max_chars: int = 4000) -> str:
#     """Truncate user/query text for safe logging (single-line friendly)."""
#     if text is None:
#         return ""
#     t = str(text).replace("\r\n", "\n").strip()
#     if len(t) <= max_chars:
#         return t
#     return t[: max_chars - 24] + "\n... [log truncated] ..."


# def configure_rag_file_logging(log_dir: str = "data/logs", filename: str = "rag_timing.log") -> None:
#     """
#     Configure project logging. This now writes to a daily file named `YYYY-MM-DD.log`
#     inside `log_dir` and appends (no truncation).

#     Notes:
#     - Each new day automatically switches to a new file even if the app stays running.
#     - Kept for backwards compatibility; prefer using ``logger.get_logger``.
#     """
#     try:
#         from logger import get_logger

#         # Configure main project logger (and avoid duplicate handlers across Streamlit reruns).
#         get_logger("yuktra_qna", level=logging.INFO, log_dir=log_dir, also_console=False)
#     except Exception:
#         # Very defensive fallback: don't crash the app just because logging couldn't be configured.
#         root = logging.getLogger("yuktra_qna")
#         if getattr(root, "_rag_file_logging_configured", False):
#             return
#         os.makedirs(log_dir, exist_ok=True)
#         path = os.path.join(log_dir, filename)
#         fh = logging.FileHandler(path, mode="a", encoding="utf-8")
#         fh.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(name)s %(message)s"))
#         root.setLevel(logging.INFO)
#         root.addHandler(fh)
#         root._rag_file_logging_configured = True  # type: ignore[attr-defined]


# def log_llm_generation_duration(
#     duration_sec: float,
#     *,
#     prompt_chars: int = 0,
#     answer_chars: int = 0,
#     used_fallback: bool = False,
# ) -> None:
#     """Call from the app after llama.cpp / HF generation finishes."""
#     logger.info(
#         "llm_generation duration_sec=%.4f prompt_chars=%d answer_chars=%d used_extractive_fallback=%s",
#         duration_sec,
#         prompt_chars,
#         answer_chars,
#         str(used_fallback),
#     )

# import numpy as np
# from pypdf import PdfReader
# from tqdm import tqdm
# import faiss  # type: ignore
# #
# SUPPORTED_TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".json", ".log"}

# # EmbeddingGemma (and similar) use asymmetric prompts for retrieval; see HF model card.
# EMBEDDING_PROMPT_STYLE_EMBEDDINGGEMMA = "embeddinggemma"
# EMBEDDING_PROMPT_STYLE_PLAIN = "plain"


# def resolve_embedding_prompt_style(config: Optional[Dict[str, Any]], embedding_model_id: str) -> str:
#     """
#     `embeddinggemma` enables official query/document prefixes at embed time.
#     Legacy vector stores omit this key — keep `plain` for queries so they match indexed vectors.
#     """
#     key = (config or {}).get("embedding_prompt_style")
#     if key in (EMBEDDING_PROMPT_STYLE_EMBEDDINGGEMMA, EMBEDDING_PROMPT_STYLE_PLAIN):
#         return str(key)
#     return EMBEDDING_PROMPT_STYLE_PLAIN


# def apply_embedding_prompt(text: str, *, role: str, style: str) -> str:
#     if style != EMBEDDING_PROMPT_STYLE_EMBEDDINGGEMMA:
#         return text
#     t = (text or "").strip()
#     if not t:
#         return text
#     if role == "query":
#         return f"task: search result | query: {t}"
#     if role == "document":
#         return f"title: none | text: {t}"
#     return text


# def is_embeddinggemma_model(model_id_or_path: str) -> bool:
#     return "embeddinggemma" in (model_id_or_path or "").lower()

# def read_text_file(path: str) -> str:
#     with open(path, "r", encoding="utf-8", errors="ignore") as f:
#         return f.read()

# def extract_text_from_pdf(path: str) -> str:
#     reader = PdfReader(path)
#     parts: List[str] = []
#     for page in reader.pages:
#         page_text = page.extract_text() or ""
#         parts.append(page_text)
#     return "\n".join(parts).strip()


# def extract_pdf_pages(path: str) -> List[Tuple[int, str]]:
#     """
#     Return per-page extracted text for PDFs so ingestion can preserve page numbers.
#     Page numbers are 1-indexed.
#     """
#     reader = PdfReader(path)
#     out: List[Tuple[int, str]] = []
#     for i, page in enumerate(reader.pages):
#         page_text = (page.extract_text() or "").strip()
#         if not page_text:
#             continue
#         out.append((i + 1, page_text))
#     return out

# def extract_text_from_docx(path: str) -> str:
#     # python-docx is an optional dependency at runtime for docx files.
#     from docx import Document  # type: ignore

#     doc = Document(path)
#     parts: List[str] = []
#     for p in doc.paragraphs:
#         if p.text and p.text.strip():
#             parts.append(p.text.strip())
#     return "\n".join(parts).strip()

# def extract_text(path: str) -> str:
#     ext = os.path.splitext(path)[1].lower()
#     if ext == ".pdf":
#         return extract_text_from_pdf(path)
#     if ext == ".docx":
#         return extract_text_from_docx(path)
#     if ext in SUPPORTED_TEXT_EXTS:
#         return read_text_file(path)
#     raise ValueError(f"Unsupported document type: {ext} ({path})")

# def chunk_text(text: str, chunk_size_chars: int, overlap_chars: int) -> List[str]:
#     """
#     Character-based chunking (offline-friendly, no tokenizers required).
#     For most RAG setups, this is "good enough" when chunk sizes are tuned.
#     """
#     if chunk_size_chars <= 0:
#         raise ValueError("chunk_size_chars must be > 0")
#     if overlap_chars < 0:
#         raise ValueError("overlap_chars must be >= 0")
#     if overlap_chars >= chunk_size_chars:
#         raise ValueError("overlap_chars must be < chunk_size_chars")

#     text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
#     if not text:
#         return []

#     # Prefer splitting on paragraph boundaries, then fall back to hard slicing.
#     paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
#     chunks: List[str] = []
#     current = ""

#     def flush():
#         nonlocal current
#         if current.strip():
#             chunks.append(current.strip())
#         current = ""

#     for para in paragraphs:
#         if len(current) + len(para) + 2 <= chunk_size_chars:
#             current = (current + "\n\n" + para).strip() if current else para
#             continue

#         # Current chunk is full; flush and start new chunk.
#         flush()
#         if len(para) <= chunk_size_chars:
#             current = para
#         else:
#             # Paragraph itself is too large; slice with overlap.
#             start = 0
#             while start < len(para):
#                 end = min(len(para), start + chunk_size_chars)
#                 piece = para[start:end]
#                 chunks.append(piece.strip())
#                 if end == len(para):
#                     break
#                 start = max(0, end - overlap_chars)

#     flush()

#     # Second pass: enforce overlap between consecutive chunks (best-effort).
#     if overlap_chars == 0 or len(chunks) <= 1:
#         return chunks
#     out: List[str] = [chunks[0]]
#     for prev, cur in zip(chunks, chunks[1:]):
#         # If overlap is needed, we can trim the start of the new chunk to simulate overlap.
#         trimmed = cur
#         if len(trimmed) > overlap_chars:
#             trimmed = trimmed  # leave as-is; the first pass already respects boundaries
#         out.append(trimmed)
#     return out

# def l2_normalize(x: np.ndarray, axis: int = -1, eps: float = 1e-12) -> np.ndarray:
#     norm = np.linalg.norm(x, axis=axis, keepdims=True)
#     return x / np.maximum(norm, eps)

# def load_llamacpp_embedding_model(
#     model_path: str,
#     *,
#     n_ctx: int = 2048,
#     n_threads: Optional[int] = None,
#     n_batch: int = 256,
#     verbose: bool = False,
# ) -> Any:
#     """
#     Load a local GGUF embedding model using llama.cpp (via llama-cpp-python).

#     Returns:
#       A `llama_cpp.Llama` instance configured for embeddings.
#     """
#     try:
#         from llama_cpp import Llama  # type: ignore
#     except Exception as e:  # pragma: no cover
#         raise RuntimeError(
#             "llama-cpp-python is required for llama.cpp embeddings. "
#             "Install it (e.g. `pip install llama-cpp-python`) and try again."
#         ) from e

#     if not model_path or not os.path.isfile(model_path):
#         raise FileNotFoundError(f"Embedding GGUF not found: {model_path}")

#     logger.info(
#         "load_llamacpp_embedding_model start path=%s n_ctx=%s n_threads=%s n_batch=%s",
#         model_path,
#         n_ctx,
#         n_threads,
#         n_batch,
#     )
#     t0 = time.perf_counter()
#     llm = Llama(
#         model_path=model_path,
#         embedding=True,
#         n_ctx=int(n_ctx),
#         n_threads=int(n_threads) if n_threads else None,
#         n_batch=int(n_batch),
#         verbose=bool(verbose),
#     )
#     logger.info(
#         "load_llamacpp_embedding_model done duration_sec=%.4f path=%s",
#         time.perf_counter() - t0,
#         model_path,
#     )
#     return llm


# def _extract_llamacpp_embeddings(resp: Any) -> np.ndarray:
#     """
#     llama-cpp-python embedding response normalizer.
#     Expected shape:
#       {"data": [{"embedding": [..]}, ...]}
#     """
#     if not isinstance(resp, dict):
#         raise RuntimeError("Unexpected llama.cpp embedding response type.")
#     data = resp.get("data")
#     if not isinstance(data, list) or not data:
#         raise RuntimeError("llama.cpp embedding response missing `data`.")
#     vecs: List[np.ndarray] = []
#     for row in data:
#         emb = (row or {}).get("embedding")
#         if not isinstance(emb, list) or not emb:
#             raise RuntimeError("llama.cpp embedding response row missing `embedding`.")
#         vecs.append(np.asarray(emb, dtype=np.float32))
#     out = np.vstack([v.reshape(1, -1) for v in vecs]).astype(np.float32)
#     return out


# def embed_texts_llamacpp(
#     texts: List[str],
#     llm: Any,
#     *,
#     batch_size: int = 16,
#     embedding_prompt_style: str = EMBEDDING_PROMPT_STYLE_PLAIN,
#     embedding_prompt_role: str = "document",
# ) -> np.ndarray:
#     t_embed = time.perf_counter()
#     total_chars = sum(len(t or "") for t in texts)
#     if not texts:
#         return np.zeros((0, 0), dtype=np.float32)

#     logger.info(
#         "embed_texts_llamacpp start num_texts=%d total_input_chars=%d batch_size=%d role=%s style=%s",
#         len(texts),
#         total_chars,
#         batch_size,
#         embedding_prompt_role,
#         embedding_prompt_style,
#     )
#     all_vecs: List[np.ndarray] = []
#     for i in tqdm(range(0, len(texts), batch_size), desc="Embedding chunks"):
#         batch_raw = texts[i : i + batch_size]
#         batch = [
#             apply_embedding_prompt(t, role=embedding_prompt_role, style=embedding_prompt_style)
#             for t in batch_raw
#         ]
#         if not hasattr(llm, "create_embedding"):
#             raise RuntimeError("llama.cpp embedding model missing `create_embedding` method.")

#         # Some llama.cpp embedding models work reliably only with single-input calls.
#         # Use per-text calls for stability.
#         batch_vecs: List[np.ndarray] = []
#         for t in batch:
#             resp = llm.create_embedding(input=t)  # type: ignore[attr-defined]
#             v = _extract_llamacpp_embeddings(resp)
#             if v.shape[0] != 1:
#                 raise RuntimeError("Expected a single embedding vector per input string.")
#             batch_vecs.append(v)
#         vecs = np.vstack(batch_vecs).astype(np.float32)
#         vecs = l2_normalize(vecs, axis=1)
#         all_vecs.append(vecs)

#     out = np.vstack(all_vecs).astype(np.float32)
#     logger.info(
#         "embed_texts_llamacpp duration_sec=%.4f num_texts=%d total_input_chars=%d prompt_role=%s style=%s",
#         time.perf_counter() - t_embed,
#         len(texts),
#         total_chars,
#         embedding_prompt_role,
#         embedding_prompt_style,
#     )
#     return out


# def save_vector_store(out_dir: str, vectors: np.ndarray, metadata: List[Dict[str, Any]], config: Dict[str, Any]) -> None:
#     logger.info(
#         "save_vector_store start out_dir=%s num_vectors=%d embedding_dim=%s index_name=%s tenant=%s",
#         out_dir,
#         int(vectors.shape[0]) if vectors.ndim == 2 else -1,
#         int(vectors.shape[1]) if vectors.ndim == 2 else -1,
#         config.get("index_name", ""),
#         config.get("tenant_name", ""),
#     )
#     os.makedirs(out_dir, exist_ok=True)
#     meta_path = os.path.join(out_dir, "metadata.json")
#     cfg_path = os.path.join(out_dir, "config.json")
#     faiss_path = os.path.join(out_dir, "index.faiss")

#     if vectors.ndim != 2:
#         raise ValueError("vectors must be 2D for FAISS export")
#     if vectors.dtype != np.float32:
#         vectors = vectors.astype(np.float32)

#     # Persist FAISS index as the primary vector store artifact.
#     dim = int(vectors.shape[1])
#     hnsw_m = 32
#     index = faiss.IndexHNSWFlat(dim, hnsw_m, faiss.METRIC_INNER_PRODUCT)
#     index.hnsw.efConstruction = 200
#     index.hnsw.efSearch = 64
#     index.add(vectors)

#     faiss.write_index(index, faiss_path)
#     with open(meta_path, "w", encoding="utf-8") as f:
#         json.dump(metadata, f, ensure_ascii=False)
#     cfg_with_faiss = dict(config)
#     cfg_with_faiss.update(
#         {
#             "vector_store_format": "faiss",
#             "faiss_index_file": os.path.basename(faiss_path),
#             "faiss_index_type": "IndexHNSWFlat",
#             "faiss_metric": "inner_product",
#             "faiss_hnsw_m": hnsw_m,
#             "faiss_ef_construction": 200,
#             "faiss_ef_search": 64,
#         }
#     )
#     with open(cfg_path, "w", encoding="utf-8") as f:
#         json.dump(cfg_with_faiss, f, ensure_ascii=False, indent=2)
#     logger.info(
#         "save_vector_store done out_dir=%s wrote_faiss=%s metadata_rows=%d",
#         out_dir,
#         os.path.basename(faiss_path),
#         len(metadata),
#     )

# def load_vector_store(
#     store_dir: str,
# ) -> Tuple[Optional[np.ndarray], Optional[Any], List[Dict[str, Any]], Dict[str, Any]]:
#     """
#     Load a vector store directory.

#     Returns:
#       (vectors, faiss_index, metadata, config)
#       Exactly one of vectors or faiss_index is non-None (FAISS-only export uses faiss_index).
#     """
#     meta_path = os.path.join(store_dir, "metadata.json")
#     cfg_path = os.path.join(store_dir, "config.json")
#     if not os.path.isdir(store_dir):
#         raise FileNotFoundError(f"Vector store directory not found: {store_dir}")
#     if not os.path.isfile(meta_path) or not os.path.isfile(cfg_path):
#         raise FileNotFoundError(
#             f"Vector store incomplete (need metadata.json + config.json): {store_dir}"
#         )

#     logger.info("load_vector_store start store_dir=%s", store_dir)

#     with open(meta_path, "r", encoding="utf-8") as f:
#         metadata = json.load(f)
#     with open(cfg_path, "r", encoding="utf-8") as f:
#         config = json.load(f)

#     config = merge_store_runtime_config(config)

#     vec_path = os.path.join(store_dir, "vectors.npy")
#     faiss_name = str(config.get("faiss_index_file") or "index.faiss")
#     faiss_path = os.path.join(store_dir, faiss_name)

#     if os.path.isfile(faiss_path):
#         index = faiss.read_index(faiss_path)
#         if hasattr(index, "hnsw") and config.get("faiss_ef_search") is not None:
#             try:
#                 index.hnsw.efSearch = int(config["faiss_ef_search"])
#             except Exception:
#                 pass
#         n = int(index.ntotal)
#         if len(metadata) != n:
#             raise ValueError(
#                 f"metadata length ({len(metadata)}) != FAISS index size ({n})."
#             )
#         logger.info(
#             "load_vector_store done backend=faiss ntotal=%d dim=%d format=%s",
#             n,
#             int(getattr(index, "d", 0) or 0),
#             config.get("vector_store_format", "faiss"),
#         )
#         return None, index, metadata, config

#     if os.path.isfile(vec_path):
#         vectors = np.load(vec_path).astype(np.float32)
#         if len(metadata) != vectors.shape[0]:
#             raise ValueError(
#                 f"metadata length ({len(metadata)}) != vectors count ({vectors.shape[0]})."
#             )
#         logger.info(
#             "load_vector_store done backend=numpy rows=%d dim=%d",
#             int(vectors.shape[0]),
#             int(vectors.shape[1]),
#         )
#         return vectors, None, metadata, config

#     raise FileNotFoundError(
#         f"No vector data found in {store_dir} (expected index.faiss or vectors.npy)."
#     )


# def topk_search_faiss(index: Any, query_vec: np.ndarray, top_k: int) -> Tuple[np.ndarray, np.ndarray]:
#     """Inner-product search (vectors expected L2-normalized). Returns (indices, scores) best-first."""
#     q = np.ascontiguousarray(query_vec.astype(np.float32).reshape(1, -1))
#     ntotal = int(index.ntotal)
#     k = min(int(top_k), max(ntotal, 0))
#     if k <= 0:
#         return np.zeros((0,), dtype=np.int64), np.zeros((0,), dtype=np.float32)
#     scores, ids = index.search(q, k)
#     row_ids = ids[0].astype(np.int64)
#     row_scores = scores[0].astype(np.float32)
#     valid = row_ids >= 0
#     return row_ids[valid], row_scores[valid]


# def _faiss_reconstruct_rows(index: Any, row_ids: np.ndarray) -> Optional[np.ndarray]:
#     """Stack reconstruct(i) for each id; None if not supported."""
#     try:
#         rows: List[np.ndarray] = []
#         for i in np.asarray(row_ids, dtype=np.int64).reshape(-1):
#             rows.append(np.asarray(index.reconstruct(int(i)), dtype=np.float32))
#         if not rows:
#             return np.zeros((0, index.d), dtype=np.float32)
#         return np.vstack(rows).astype(np.float32)
#     except Exception:
#         return None


# def topk_search(vectors: np.ndarray, query_vec: np.ndarray, top_k: int) -> Tuple[np.ndarray, np.ndarray]:
#     """
#     Cosine similarity via dot product (vectors are expected normalized).
#     Returns (indices, scores) sorted by descending score.
#     """
#     if vectors.ndim != 2:
#         raise ValueError("vectors must be 2D")
#     if query_vec.ndim == 2 and query_vec.shape[0] == 1:
#         query_vec = query_vec[0]
#     if query_vec.ndim != 1:
#         raise ValueError("query_vec must be 1D")

#     scores = vectors @ query_vec.astype(np.float32)  # [N]
#     if top_k >= len(scores):
#         idx = np.argsort(-scores)
#     else:
#         # partial sort for speed, then fully sort top-k slice
#         idx = np.argpartition(-scores, top_k)[:top_k]
#         idx = idx[np.argsort(-scores[idx])]
#     return idx, scores[idx]


# def mmr_select_indices(
#     pool_row_indices: np.ndarray,
#     vectors: np.ndarray,
#     query_vec: np.ndarray,
#     k_out: int,
#     lambda_mult: float = 0.62,
# ) -> Tuple[np.ndarray, np.ndarray]:
#     """
#     Maximal Marginal Relevance on a candidate pool (row indices into `vectors`).
#     Returns (ordered_row_indices, scores[dot with query]) sorted by semantic score descending.
#     """
#     pool_row_indices = np.asarray(pool_row_indices, dtype=np.int64).reshape(-1)
#     if pool_row_indices.size == 0:
#         return pool_row_indices, np.zeros((0,), dtype=np.float32)
#     q = query_vec.astype(np.float32).reshape(-1)
#     if pool_row_indices.size <= k_out:
#         sel = pool_row_indices
#         sc = vectors[sel] @ q
#         order = np.argsort(-sc)
#         return sel[order], sc[order]

#     cand_vecs = vectors[pool_row_indices].astype(np.float32)
#     sim_q = cand_vecs @ q
#     sim_mx = cand_vecs @ cand_vecs.T
#     m = cand_vecs.shape[0]
#     k_out = min(k_out, m)

#     selected_local: List[int] = []
#     selected_set = set()
#     first = int(np.argmax(sim_q))
#     selected_local.append(first)
#     selected_set.add(first)

#     while len(selected_local) < k_out:
#         best_j = -1
#         best_mmr = -1e9
#         for j in range(m):
#             if j in selected_set:
#                 continue
#             max_sim_sel = max(float(sim_mx[j, i]) for i in selected_local)
#             mmr = lambda_mult * float(sim_q[j]) - (1.0 - lambda_mult) * max_sim_sel
#             if mmr > best_mmr:
#                 best_mmr = mmr
#                 best_j = j
#         selected_local.append(best_j)
#         selected_set.add(best_j)

#     sel_rows = pool_row_indices[np.array(selected_local, dtype=np.int64)]
#     sc = vectors[sel_rows] @ q
#     order = np.argsort(-sc)
#     return sel_rows[order], sc[order]


# def pack_context_from_reranked(
#     reranked: List[Tuple[Dict[str, Any], float]],
#     max_chars: int,
#     min_slice_chars: int = 280,
#     max_chunks: Optional[int] = None,
# ) -> List[Dict[str, Any]]:
#     """
#     Fill the LLM context budget in rerank order, but skip oversized chunks when the
#     budget is low so later smaller (still relevant) chunks can be included.
#     """
#     if max_chars <= 0:
#         return []
#     if max_chunks is not None and max_chunks > 0:
#         reranked = reranked[:max_chunks]
#     out: List[Dict[str, Any]] = []
#     remaining = max_chars
#     for ch, _ in reranked:
#         text = strip_encoded_payload_noise(str(ch.get("text", "")))
#         if not text.strip():
#             continue
#         ch = dict(ch)
#         ch["text"] = text
#         if len(text) <= remaining:
#             out.append(ch)
#             remaining -= len(text)
#             if remaining <= 0:
#                 break
#             continue
#         if remaining >= min_slice_chars:
#             partial = dict(ch)
#             partial["text"] = text[:remaining]
#             out.append(partial)
#             break
#     return out


# def _parse_page_number(value: Any) -> Optional[int]:
#     if isinstance(value, int):
#         return value if value > 0 else None
#     s = str(value or "").strip()
#     if not s:
#         return None
#     m = re.search(r"\d+", s)
#     if not m:
#         return None
#     try:
#         p = int(m.group(0))
#         return p if p > 0 else None
#     except Exception:
#         return None


# def _parse_chunk_index(value: Any) -> Optional[int]:
#     if value is None:
#         return None
#     try:
#         ci = int(value)
#         return ci if ci >= 0 else None
#     except (TypeError, ValueError):
#         return None


# def expand_reranked_with_doc_chunk_neighbors(
#     reranked: List[Tuple[Dict[str, Any], float]],
#     metadata: List[Dict[str, Any]],
#     *,
#     neighbor_before: int = 5,
#     neighbor_after: int = 1,
# ) -> List[Tuple[Dict[str, Any], float]]:
#     """
#     After hybrid rerank and score filtering, widen each remaining hit along the same
#     document using chunk_index neighbors in a directional window
#     (current-neighbor_before..current+neighbor_after).
#     Seed chunks stay first in original order; neighbors append sorted by (doc_name, chunk_index).
#     """
#     if not reranked or (neighbor_before <= 0 and neighbor_after <= 0):
#         return list(reranked)

#     by_doc_ci: Dict[Tuple[str, int], Dict[str, Any]] = {}
#     for m in metadata:
#         doc = str(m.get("doc_name", "") or "")
#         ci = _parse_chunk_index(m.get("chunk_index"))
#         if not doc or ci is None:
#             continue
#         by_doc_ci[(doc, ci)] = m

#     def _row_key(ch: Dict[str, Any]) -> Tuple[str, int, int]:
#         doc = str(ch.get("doc_name", "") or "")
#         vec_id = int(ch.get("vector_id", -1) or -1)
#         ci = _parse_chunk_index(ch.get("chunk_index"))
#         return (doc, vec_id, ci if ci is not None else -1)

#     out: List[Tuple[Dict[str, Any], float]] = []
#     seen: set[Tuple[str, int, int]] = set()
#     for ch, sc in reranked:
#         rk = _row_key(ch)
#         if rk in seen:
#             continue
#         seen.add(rk)
#         out.append((ch, float(sc)))

#     want_by_doc: Dict[str, set[int]] = {}
#     for ch, _ in reranked:
#         doc = str(ch.get("doc_name", "") or "")
#         ci = _parse_chunk_index(ch.get("chunk_index"))
#         if not doc or ci is None:
#             continue
#         s = want_by_doc.setdefault(doc, set())
#         for d in range(-max(0, int(neighbor_before)), max(0, int(neighbor_after)) + 1):
#             nj = ci + d
#             if nj >= 0:
#                 s.add(nj)

#     extra: List[Tuple[Dict[str, Any], float]] = []
#     for doc in sorted(want_by_doc.keys()):
#         for ci in sorted(want_by_doc[doc]):
#             m = by_doc_ci.get((doc, ci))
#             if m is None:
#                 continue
#             rk = _row_key(m)
#             if rk in seen:
#                 continue
#             seen.add(rk)
#             extra.append((m, 0.0))

#     if extra:
#         logger.info(
#             "expand_reranked_with_doc_chunk_neighbors seeds=%d added_neighbors=%d before=%d after=%d",
#             len(out),
#             len(extra),
#             max(0, int(neighbor_before)),
#             max(0, int(neighbor_after)),
#         )
#     return out + extra


# def _preview_text_for_log(text: str, *, head: int = 140, tail: int = 90) -> str:
#     t = " ".join((text or "").split())
#     if not t:
#         return ""
#     if len(t) <= head + tail + 5:
#         return t
#     return f"{t[:head]} … {t[-tail:]}"


# def _format_chunk_for_retrieval_log(ch: Dict[str, Any]) -> str:
#     doc = str(ch.get("doc_name", "") or "")
#     vec_id = ch.get("vector_id", "?")
#     chunk_idx = ch.get("chunk_index", "?")
#     page = ch.get("page_number", "?")
#     sec = str(ch.get("section_path_str", "") or "")
#     sec = " ".join(sec.split())
#     if len(sec) > 80:
#         sec = sec[:77] + "…"
#     txt = _preview_text_for_log(str(ch.get("text", "") or ""))
#     return (
#         f"doc={doc} page={page} chunk={chunk_idx} vec={vec_id}"
#         + (f" section={sec}" if sec else "")
#         + (f" text=\"{txt}\"" if txt else " text=\"\"")
#     )


# def _log_retrieval_chunk_stage(
#     label: str,
#     rows: List[Tuple[Dict[str, Any], float]],
#     *,
#     max_items: int = 256,
# ) -> None:
#     """One line per chunk for a retrieval phase (initial / filtered / remaining, etc.)."""
#     n = len(rows)
#     logger.info("retrieve_rag_pipeline %s count=%d", label, n)
#     cap = min(n, max(0, int(max_items)))
#     for i in range(cap):
#         ch, sc = rows[i]
#         logger.info(
#             "retrieve_rag_pipeline %s #%d score=%.4f %s",
#             label,
#             i + 1,
#             float(sc),
#             _format_chunk_for_retrieval_log(ch),
#         )
#     if n > cap:
#         logger.info("retrieve_rag_pipeline %s ... %d more omitted (max_items=%d)", label, n - cap, max_items)


# def pack_top_doc_neighbor_pages_context(
#     reranked: List[Tuple[Dict[str, Any], float]],
#     metadata: List[Dict[str, Any]],
#     *,
#     max_chars: int,
#     page_window: int = 1,
#     max_chunks: Optional[int] = None,
#     question: str = "",
# ) -> List[Dict[str, Any]]:
#     """
#     Build context from the top document only: every reranked row for `doc_name` of
#     the first hit, in list order (hybrid-ranked seeds followed by same-doc neighbors
#     from `expand_reranked_with_doc_chunk_neighbors`), capped to `max_chars` and optional
#     `max_chunks`.

#     `metadata` and `page_window` are kept for backward compatibility with call sites;
#     neighbor expansion is done in `retrieve_rag_pipeline` before this runs.
#     """
#     if max_chars <= 0 or not reranked:
#         return []

#     top_doc = str(reranked[0][0].get("doc_name", "") or "")
#     if not top_doc:
#         return pack_context_from_reranked(reranked, max_chars, max_chunks=max_chunks)

#     ordered: List[Dict[str, Any]] = []
#     seen_keys: set[Tuple[str, int, int]] = set()

#     def _k(ch: Dict[str, Any]) -> Tuple[str, int, int]:
#         doc = str(ch.get("doc_name", "") or "")
#         chunk_idx = int(ch.get("chunk_index", -1) or -1)
#         vec_id = int(ch.get("vector_id", -1) or -1)
#         return (doc, vec_id, chunk_idx)

#     # Top document only: collect unique chunks first.
#     for ch, _ in reranked:
#         if str(ch.get("doc_name", "") or "") != top_doc:
#             continue
#         key = _k(ch)
#         if key in seen_keys:
#             continue
#         seen_keys.add(key)
#         ordered.append(ch)

#     if not ordered:
#         return pack_context_from_reranked(reranked, max_chars, max_chunks=max_chunks)

#     # Deterministic chronological/manual flow for the LLM:
#     # page asc -> chunk_index asc -> vector_id asc.
#     # This avoids seed-first ordering like 19,18,20... and keeps context as 18,19,20,21...
#     def _llm_context_order_key(ch: Dict[str, Any]) -> Tuple[int, int, int]:
#         p = _parse_page_number(ch.get("page_number"))
#         ci = _parse_chunk_index(ch.get("chunk_index"))
#         try:
#             vid = int(ch.get("vector_id", -1) or -1)
#         except Exception:
#             vid = -1
#         return (
#             p if p is not None else 10**9,
#             ci if ci is not None else 10**9,
#             vid,
#         )

#     ordered.sort(key=_llm_context_order_key)

#     out: List[Dict[str, Any]] = []
#     remaining = max_chars
#     for ch in ordered:
#         text = strip_encoded_payload_noise(str(ch.get("text", "") or ""))
#         if not text.strip():
#             continue
#         ch = dict(ch)
#         ch["text"] = text
#         if max_chunks is not None and max_chunks > 0 and len(out) >= max_chunks:
#             break
#         if len(text) <= remaining:
#             out.append(ch)
#             remaining -= len(text)
#             if remaining <= 0:
#                 break
#             continue
#         if remaining >= 280:
#             partial = dict(ch)
#             partial["text"] = text[:remaining]
#             out.append(partial)
#         break
#     return out


# def retrieval_query_variants(question: str) -> List[str]:
#     """
#     Lightweight query expansion for embedding fusion (no extra LLM call).
#     Improves recall for specs, troubleshooting, and safety questions.
#     """
#     q = (question or "").strip()
#     if not q:
#         return ["equipment manual"]
#     variants = [q]
#     ql = q.lower()

#     if any(
#         k in ql
#         for k in (
#             "voltage",
#             "volt",
#             "vac",
#             "vdc",
#             "electrical",
#             "power supply",
#             "line voltage",
#             "mains",
#             "frequency",
#             "hz",
#             "amp",
#             "current",
#         )
#     ):
#         variants.append(
#             f"{q} rated input voltage AC DC electrical specifications power consumption frequency"
#         )

#     if any(
#         k in ql
#         for k in (
#             "troubleshoot",
#             "troubleshooting",
#             "not light",
#             "does not light",
#             "indicator",
#             "power indicator",
#             "lamp",
#             "fault",
#             "not working",
#             "doesn't work",
#             "won't turn",
#         )
#     ):
#         variants.append(
#             f"{q} troubleshooting diagnostic fault repair power connection indicator lamp supply"
#         )

#     if any(
#         k in ql
#         for k in ("safe", "safety", "hazard", "precaution", "warning", "ppe", "protective", "eye")
#     ):
#         variants.append(f"{q} safety hazard warning precaution operation protective equipment")

#     if any(k in ql for k in ("weight", "kg", "dimension", "size", "mass")):
#         variants.append(f"{q} physical specifications weight dimensions")

#     if any(
#         k in ql
#         for k in (
#             "calibrat",
#             "verification",
#             "qualification",
#             "validation protocol",
#             " as-found",
#             "as-left",
#             "iq ",
#             "oq ",
#             "pq ",
#         )
#     ):
#         variants.append(
#             f"{q} calibration verification qualification adjustment tolerance specification check procedure"
#         )

#     if any(
#         k in ql
#         for k in (
#             "temperature",
#             "humidity",
#             "dew point",
#             "relative humidity",
#             "rh ",
#             "pressure",
#             " psi",
#             " bar",
#             " kpa",
#             " mpa",
#         )
#     ):
#         variants.append(f"{q} environmental operating conditions range limits specifications")

#     if any(
#         k in ql
#         for k in (
#             "chromatograph",
#             "hplc",
#             "uhplc",
#             "lc-ms",
#             "gc-ms",
#             "2d-lc",
#             "2d lc",
#             "two-dimensional",
#             "heart-cut",
#             "heart cut",
#             "heart-cutting",
#             "multiple heart",
#             "comprehensive 2d",
#             "lc x lc",
#             "lclc",
#         )
#     ):
#         variants.append(
#             f"{q} separation column valve sample loop injection modulation multidimensional chromatography method "
#             "principle backward reverse order contamination flush gradient parking deck peak lost occupied loop analysis"
#         )

#     if any(
#         k in ql
#         for k in (
#             "what is",
#             "define",
#             "definition",
#             "principle",
#             "principles",
#             "overview",
#             "purpose of",
#         )
#     ):
#         variants.append(f"{q} characteristics description notes operation overview")

#     seen: set[str] = set()
#     out: List[str] = []
#     for v in variants:
#         k = v.lower()
#         if k not in seen:
#             seen.add(k)
#             out.append(v)
#     return out


# BM25_LEXICAL_QUERY_MAX_CHARS = 900


# def build_bm25_lexical_query(question: str, max_chars: int = BM25_LEXICAL_QUERY_MAX_CHARS) -> str:
#     """
#     BM25 over the initial retrieval pool: original question plus the first expansion variant
#     (from retrieval_query_variants), capped for tokenizer/BM25 stability.
#     """
#     variants = retrieval_query_variants(question)
#     base = (variants[0] if variants else (question or "")).strip()
#     if len(variants) <= 1:
#         return base[:max_chars] if max_chars > 0 else base
#     extra = variants[1].strip()
#     combined = f"{base} {extra}".strip()
#     if max_chars <= 0 or len(combined) <= max_chars:
#         return combined
#     return combined[:max_chars]


# def embed_fused_query_for_retrieval(
#     question: str,
#     tokenizer: Any,
#     model: Any,
#     device: str,
#     max_length: int,
#     embedding_prompt_style: str,
#     batch_size: int = 8,
# ) -> np.ndarray:
#     """
#     Embed query variants and L2-normalize the mean vector (better recall than a single embed).
#     """
#     t0 = time.perf_counter()
#     q_preview = pipeline_log_preview(question, max_chars=2000)
#     logger.info(
#         "embed_fused_query_for_retrieval start query_chars=%d style=%s max_length=%d device=%s query=%r",
#         len(question or ""),
#         embedding_prompt_style,
#         max_length,
#         device,
#         q_preview,
#     )
#     variants = retrieval_query_variants(question)
#     bs = max(1, min(batch_size, len(variants)))
#     backend = "llamacpp"
#     logger.info(
#         "embed_fused_query_for_retrieval variants_ready count=%d batch_size=%d backend=%s",
#         len(variants),
#         bs,
#         backend,
#     )
#     if not hasattr(model, "create_embedding"):
#         raise RuntimeError(
#             "Embedding backend must be llama.cpp (expected `create_embedding` on model)."
#         )
#     vecs = embed_texts_llamacpp(
#         variants,
#         model,
#         batch_size=bs,
#         embedding_prompt_style=embedding_prompt_style,
#         embedding_prompt_role="query",
#     )
#     fused = np.mean(vecs.astype(np.float32), axis=0)
#     out = l2_normalize(fused.reshape(1, -1), axis=1)[0]
#     logger.info(
#         "embed_fused_query_for_retrieval done duration_sec=%.4f dim=%d backend=%s",
#         time.perf_counter() - t0,
#         int(out.shape[0]),
#         backend,
#     )
#     return out


# def _minmax_norm_1d(x: np.ndarray, eps: float = 1e-8) -> np.ndarray:
#     x = np.asarray(x, dtype=np.float32)
#     lo = float(x.min())
#     hi = float(x.max())
#     if hi - lo < eps:
#         return np.ones_like(x, dtype=np.float32)
#     return ((x - lo) / (hi - lo + eps)).astype(np.float32)


# def _tokenize_bm25(s: str) -> List[str]:
#     return re.findall(r"[a-zA-Z0-9]+", (s or "").lower())


# def bm25_scores_for_pool(
#     query: str,
#     doc_texts: List[str],
#     k1: float = 1.2,
#     b: float = 0.75,
# ) -> np.ndarray:
#     """BM25 over the current candidate pool only (no external index)."""
#     N = len(doc_texts)
#     if N == 0:
#         return np.zeros(0, dtype=np.float32)
#     tokenized_docs = [_tokenize_bm25(t) for t in doc_texts]
#     doc_lens = np.array([max(1, len(td)) for td in tokenized_docs], dtype=np.float32)
#     avgdl = float(np.mean(doc_lens)) + 1e-6

#     df: Dict[str, int] = {}
#     for td in tokenized_docs:
#         for t in set(td):
#             df[t] = df.get(t, 0) + 1

#     q_terms = _tokenize_bm25(query)
#     if not q_terms:
#         return np.zeros(N, dtype=np.float32)

#     scores = np.zeros(N, dtype=np.float32)
#     for i, td in enumerate(tokenized_docs):
#         tf: Dict[str, int] = {}
#         for t in td:
#             tf[t] = tf.get(t, 0) + 1
#         dl = doc_lens[i]
#         s = 0.0
#         for qt in q_terms:
#             if qt not in tf:
#                 continue
#             dfi = df.get(qt, 0)
#             idf = float(np.log((N - dfi + 0.5) / (dfi + 0.5) + 1.0))
#             f = float(tf[qt])
#             denom = f + k1 * (1.0 - b + b * (dl / avgdl))
#             s += idf * (f * (k1 + 1.0)) / (denom + 1e-8)
#         scores[i] = s
#     return scores


# _BOILERPLATE_NOISE = re.compile(
#     r"iso\s*9001|quality\s+management\s+system|about\s+this\s+document|"
#     r"desktop\s+computer|minimum\s+requirement|ergonomic|keyboard\s+and\s+mouse|"
#     r"limitations?\s+of\s+warranty|registered\s+trademarks",
#     re.I,
# )
# _TECH_IN_CHUNK = re.compile(
#     r"voltage|vac|vdc|troubleshoot|indicator|12\s*v|24\s*v|power\s+supply|"
#     r"lamp\s+module|fuse|wiring|specification|rated\s+input|green|wall\s+outlet|figure",
#     re.I,
# )
# _ENCODED_PAYLOAD_RE = re.compile(r"\b[A-Za-z0-9+/]{96,}={0,2}\b")


# def strip_encoded_payload_noise(text: str) -> str:
#     """Remove base64-like image/OCR payloads that can leak into text chunks."""
#     t = str(text or "")
#     if not t:
#         return ""
#     t = _ENCODED_PAYLOAD_RE.sub("", t)
#     t = re.sub(r"[ \t]{2,}", " ", t)
#     t = re.sub(r"\n{3,}", "\n\n", t)
#     return t.strip()


# def encoded_payload_noise_multiplier(chunk_text: str) -> float:
#     """Strongly downrank chunks dominated by encoded image bytes instead of readable text."""
#     t = str(chunk_text or "")
#     if len(t) < 220:
#         return 1.0
#     encoded_chars = sum(len(m.group(0)) for m in _ENCODED_PAYLOAD_RE.finditer(t))
#     if encoded_chars < 160:
#         return 1.0
#     ratio = encoded_chars / max(1, len(t))
#     if ratio >= 0.55:
#         return 0.18
#     if ratio >= 0.30:
#         return 0.45
#     return 1.0


# def boilerplate_noise_multiplier(chunk_text: str) -> float:
#     """Downweight generic intro / ISO / PC-requirement chunks that lack substantive technical text."""
#     if not chunk_text:
#         return 1.0
#     head = chunk_text[:12000]
#     if _BOILERPLATE_NOISE.search(head) and not _TECH_IN_CHUNK.search(head):
#         return 0.32
#     return 1.0


# _RETRIEVAL_Q_STOP = frozenset(
#     {
#         "what",
#         "the",
#         "is",
#         "are",
#         "was",
#         "were",
#         "does",
#         "did",
#         "do",
#         "how",
#         "when",
#         "where",
#         "why",
#         "which",
#         "who",
#         "this",
#         "that",
#         "for",
#         "with",
#         "from",
#         "have",
#         "has",
#         "should",
#         "would",
#         "could",
#         "can",
#         "will",
#         "been",
#         "being",
#         "into",
#         "about",
#         "your",
#         "any",
#         "all",
#     }
# )


# def query_term_coverage_multiplier(chunk_text: str, question: str) -> float:
#     """Boost chunks that contain more of the user's content words (lexical grounding)."""
#     if not chunk_text or not (question or "").strip():
#         return 1.0
#     qts = [
#         t
#         for t in re.findall(r"[a-zA-Z]{3,}", question.lower())
#         if t not in _RETRIEVAL_Q_STOP
#     ]
#     if len(qts) < 2:
#         return 1.0
#     blob = chunk_text.lower()
#     hits = sum(1 for t in set(qts) if t in blob)
#     ratio = hits / max(1, len(set(qts)))
#     return float(0.78 + 0.50 * min(1.0, ratio))


# def score_filter_lexical_fallback(chunk_text: str, question: str) -> Tuple[int, float]:
#     """
#     Decide whether a below-threshold rerank hit is still too lexically relevant to drop.

#     This catches narrow manual answers such as "Reset alarm: when there is abnormality..."
#     where the semantic score can fall below a global threshold, but the exact user terms
#     are present in the chunk.
#     """
#     if not chunk_text or not (question or "").strip():
#         return 0, 0.0
#     q_terms = {
#         t
#         for t in re.findall(r"[a-zA-Z]{3,}", question.lower())
#         if t not in _RETRIEVAL_Q_STOP
#     }
#     if not q_terms:
#         return 0, 0.0
#     blob = chunk_text.lower()
#     hits = sum(1 for t in q_terms if t in blob)
#     coverage = float(hits) / float(max(1, len(q_terms)))
#     return hits, coverage


# def thin_chunk_multiplier(chunk_text: str) -> float:
#     """Slightly downrank very short chunks (often headers/noise)."""
#     L = len((chunk_text or "").strip())
#     if L < 80:
#         return 0.75
#     if L < 180:
#         return 0.9
#     return 1.0


# def _ranks_descending_scores(scores: np.ndarray) -> np.ndarray:
#     """rank[j] = 0 for best (highest score), 1 for second, ..."""
#     scores = np.asarray(scores, dtype=np.float32)
#     order = np.argsort(-scores)
#     ranks = np.empty(len(scores), dtype=np.float32)
#     ranks[order] = np.arange(len(scores), dtype=np.float32)
#     return ranks


# def reciprocal_rank_fusion_score(dense_scores: np.ndarray, bm25_scores: np.ndarray, k: float = 58.0) -> np.ndarray:
#     """
#     RRF combines two rankers without fragile score scaling (Cormack et al. style).
#     """
#     rd = _ranks_descending_scores(dense_scores)
#     rb = _ranks_descending_scores(bm25_scores)
#     return (1.0 / (k + rd) + 1.0 / (k + rb)).astype(np.float32)


# def ensure_non_empty_capped_context(
#     capped: List[Dict[str, Any]],
#     retrieved: List[Dict[str, Any]],
#     max_context_chars: int,
# ) -> List[Dict[str, Any]]:
#     """
#     If packing skipped everything (e.g. oversized chunks vs tight budget), still pass the
#     top retrieved chunk so the LLM is not sent empty context.
#     """
#     if capped:
#         return capped
#     if not retrieved:
#         return []
#     ch = dict(retrieved[0])
#     t = str(ch.get("text", ""))
#     if not t.strip():
#         return []
#     budget = max_context_chars if max_context_chars > 0 else 8001
#     limit = min(len(t), budget)
#     ch["text"] = t[:limit]
#     return [ch]


# def retrieve_rag_pipeline(
#     question: str,
#     query_vec: np.ndarray,
#     vectors: Optional[np.ndarray],
#     metadata: List[Dict[str, Any]],
#     *,
#     faiss_index: Optional[Any] = None,
#     top_k: int,
#     retrieval_pool_k: int,
#     mmr_k: int,
#     mmr_lambda: float = 0.62,
#     max_context_chars: int,
#     restrict_top_document: bool = False,
#     bm25_weight: float = 0.6,
#     initial_pool_multiplier: int = 2,
#     max_chunks_for_prompt: int = 3,
#     rrf_linear_blend: float = 0.25,
#     hybrid_alpha_semantic: float = 0.4,
#     min_hybrid_rerank_score: float = 0.0,
#     top_doc_page_window: int = 2,
#     top_doc_chunk_neighbor_radius: int = 6,
#     top_doc_chunk_neighbors_before: Optional[int] = None,
#     top_doc_chunk_neighbors_after: Optional[int] = None,
# ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Tuple[Dict[str, Any], float]]]:
#     """
#     Returns:
#       - retrieved: chunks after hybrid rerank, score threshold, and same-doc chunk_index
#         ±N neighbor expansion (via metadata, default N=6).
#       - capped_for_prompt: packed to max_context_chars (reduces noise)
#       - reranked: (chunk, score) pairs after threshold and neighbor expansion; seeds keep
#         hybrid scores, neighbors use 0.0.

#     Dense + BM25: min-max linear blend plus Reciprocal Rank Fusion (RRF) for robustness; then
#     boilerplate / coverage / thin-chunk multipliers. BM25 uses the original question plus the
#     first retrieval expansion variant (capped at BM25_LEXICAL_QUERY_MAX_CHARS).
#     """
#     t_pipe = time.perf_counter()
#     n = len(metadata)
#     if n == 0:
#         logger.info(
#             "retrieve_rag_pipeline empty_store duration_sec=%.4f query_preview=%r",
#             time.perf_counter() - t_pipe,
#             pipeline_log_preview(question, max_chars=800),
#         )
#         return [], [], []

#     if (vectors is None) == (faiss_index is None):
#         raise ValueError("retrieve_rag_pipeline: set exactly one of vectors or faiss_index")

#     pool_k = min(n, max(top_k, retrieval_pool_k))
#     mult = max(1, int(initial_pool_multiplier))
#     if n > 40:
#         mult = max(mult, 3)
#     initial_k = min(n, max(pool_k * mult, pool_k + min(32, n // 4)))

#     backend = "faiss" if faiss_index is not None else "numpy"
#     q_log = pipeline_log_preview(question, max_chars=3500)
#     logger.info(
#         "retrieve_rag_pipeline phase=init query_chars=%d backend=%s corpus_chunks=%d "
#         "top_k=%d retrieval_pool_k=%d mmr_k=%d initial_k=%d pool_k=%d max_context_chars=%d "
#         "restrict_top_doc=%s query=%r",
#         len(question or ""),
#         backend,
#         n,
#         top_k,
#         retrieval_pool_k,
#         mmr_k,
#         initial_k,
#         pool_k,
#         max_context_chars,
#         restrict_top_document,
#         q_log,
#     )

#     t_dense = time.perf_counter()
#     if faiss_index is not None:
#         idx, dense_scores = topk_search_faiss(faiss_index, query_vec, initial_k)
#     else:
#         assert vectors is not None
#         idx, dense_scores = topk_search(vectors, query_vec, initial_k)
#     texts = [str(metadata[int(i)].get("text", "")) for i in idx]
#     logger.info(
#         "retrieve_rag_pipeline phase=dense_search duration_sec=%.4f initial_k=%d returned=%d dense_min=%.4f dense_max=%.4f",
#         time.perf_counter() - t_dense,
#         initial_k,
#         len(idx),
#         float(np.min(dense_scores)) if len(dense_scores) else 0.0,
#         float(np.max(dense_scores)) if len(dense_scores) else 0.0,
#     )

#     t_fuse = time.perf_counter()
#     bm25_query = build_bm25_lexical_query(question)
#     bm25_s = bm25_scores_for_pool(bm25_query, texts)
#     dn = _minmax_norm_1d(dense_scores)
#     bn = _minmax_norm_1d(bm25_s)
#     w_bm25 = float(np.clip(bm25_weight, 0.0, 0.95))
#     linear = (1.0 - w_bm25) * dn + w_bm25 * bn
#     rrf = reciprocal_rank_fusion_score(dense_scores, bm25_s)
#     rrf_n = _minmax_norm_1d(rrf)
#     blend = float(np.clip(rrf_linear_blend, 0.0, 0.95))
#     fused = (1.0 - blend) * linear + blend * rrf_n

#     for j in range(len(idx)):
#         fused[j] *= boilerplate_noise_multiplier(texts[j])
#         fused[j] *= encoded_payload_noise_multiplier(texts[j])
#         fused[j] *= query_term_coverage_multiplier(texts[j], question)
#         fused[j] *= thin_chunk_multiplier(texts[j])

#     order = np.argsort(-fused)
#     idx = idx[order]
#     fused = fused[order]
#     idx = idx[:pool_k].copy()
#     fused = fused[:pool_k].copy()
#     logger.info(
#         "retrieve_rag_pipeline phase=hybrid_fusion duration_sec=%.4f pool_k=%d bm25_weight=%.2f rrf_blend=%.2f bm25_query_chars=%d",
#         time.perf_counter() - t_fuse,
#         pool_k,
#         w_bm25,
#         blend,
#         len(bm25_query),
#     )

#     t_mmr = time.perf_counter()
#     mmr_take = min(mmr_k, pool_k)
#     if mmr_take >= pool_k:
#         idx_sel = idx
#         scores_sel = fused
#     else:
#         if faiss_index is not None:
#             cand_matrix = _faiss_reconstruct_rows(faiss_index, idx)
#             if cand_matrix is None or cand_matrix.shape[0] != len(idx):
#                 idx_sel = idx[:mmr_take]
#                 scores_sel = fused[:mmr_take].copy()
#             else:
#                 local = np.arange(len(idx), dtype=np.int64)
#                 idx_local_sel, _ = mmr_select_indices(
#                     local, cand_matrix, query_vec, mmr_take, lambda_mult=mmr_lambda
#                 )
#                 idx_sel = idx[idx_local_sel.astype(np.int64)]
#                 lut = {int(idx[j]): float(fused[j]) for j in range(len(idx))}
#                 scores_sel = np.array([lut[int(i)] for i in idx_sel], dtype=np.float32)
#         else:
#             assert vectors is not None
#             idx_sel, _ = mmr_select_indices(idx, vectors, query_vec, mmr_take, lambda_mult=mmr_lambda)
#             lut = {int(idx[j]): float(fused[j]) for j in range(len(idx))}
#             scores_sel = np.array([lut[int(i)] for i in idx_sel], dtype=np.float32)

#     logger.info(
#         "retrieve_rag_pipeline phase=mmr duration_sec=%.4f mmr_take=%d candidate_rows=%d (pool_before_mmr=%d)",
#         time.perf_counter() - t_mmr,
#         mmr_take,
#         len(idx_sel),
#         pool_k,
#     )

#     candidates = [metadata[int(i)] for i in idx_sel]
#     candidate_scores = np.asarray(scores_sel, dtype=np.float32)
#     _log_retrieval_chunk_stage(
#         "initial_chunks",
#         [(c, float(candidate_scores[j])) for j, c in enumerate(candidates)],
#     )

#     alpha_sem = float(np.clip(hybrid_alpha_semantic, 0.05, 0.95))
#     beta_lex = 1.0 - alpha_sem

#     t_rerank = time.perf_counter()
#     reranked = hybrid_rerank(
#         question,
#         candidate_chunks=candidates,
#         candidate_scores=candidate_scores,
#         top_k=top_k,
#         alpha_semantic=alpha_sem,
#         beta_lexical=beta_lex,
#     )
#     if reranked:
#         logger.info(
#             "retrieve_rag_pipeline phase=hybrid_rerank duration_sec=%.4f rows=%d",
#             time.perf_counter() - t_rerank,
#             len(reranked),
#         )
#         for i, (ch, sc) in enumerate(reranked, start=1):
#             logger.info(
#                 "retrieve_rag_pipeline hybrid_rerank #%d doc=%s chunk_index=%s vec=%s score=%.4f",
#                 i,
#                 str(ch.get("doc_name", "?"))[:120],
#                 ch.get("chunk_index", "?"),
#                 ch.get("vector_id", "?"),
#                 float(sc),
#             )
#     else:
#         logger.info(
#             "retrieve_rag_pipeline phase=hybrid_rerank duration_sec=%.4f rows=0 (empty)",
#             time.perf_counter() - t_rerank,
#         )
#     t_post_filter = time.perf_counter()
#     if restrict_top_document:
#         reranked = enforce_single_doc_scope(question, reranked, top_k=top_k)

#     if restrict_top_document and reranked:
#         top_doc = str(reranked[0][0].get("doc_name", "") or "")
#         if top_doc:
#             reranked = [(ch, s) for ch, s in reranked if str(ch.get("doc_name", "")) == top_doc][:top_k]

#     min_hs = float(np.clip(min_hybrid_rerank_score, 0.0, 1.0))
#     if min_hs > 0.0 and reranked:
#         before = len(reranked)
#         dropped_hs = [(ch, s) for ch, s in reranked if float(s) < min_hs]
#         filtered_hs = [(ch, s) for ch, s in reranked if float(s) >= min_hs]
#         if dropped_hs:
#             _log_retrieval_chunk_stage("score_filter_dropped_chunks", dropped_hs)
#         if filtered_hs:
#             reranked = filtered_hs
#             logger.info(
#                 "retrieve_rag_pipeline phase=score_threshold min=%.3f kept=%d dropped=%d",
#                 min_hs,
#                 len(reranked),
#                 before - len(reranked),
#             )
#         else:
#             lexical_fb: List[Tuple[Dict[str, Any], float]] = []
#             for ch, s in dropped_hs:
#                 hits, coverage = score_filter_lexical_fallback(str(ch.get("text", "") or ""), question)
#                 if hits >= 2 or coverage >= 0.45:
#                     lexical_fb.append((ch, s))
#                     logger.info(
#                         "retrieve_rag_pipeline phase=score_threshold lexical_fallback_candidate "
#                         "score=%.4f hits=%d coverage=%.2f %s",
#                         float(s),
#                         hits,
#                         coverage,
#                         _format_chunk_for_retrieval_log(ch),
#                     )
#             if lexical_fb:
#                 lexical_fb.sort(
#                     key=lambda item: (
#                         score_filter_lexical_fallback(str(item[0].get("text", "") or ""), question)[1],
#                         score_filter_lexical_fallback(str(item[0].get("text", "") or ""), question)[0],
#                         float(item[1]),
#                     ),
#                     reverse=True,
#                 )
#                 reranked = lexical_fb[: max(1, min(top_k, 3))]
#                 logger.info(
#                     "retrieve_rag_pipeline phase=score_threshold min=%.3f kept=%d dropped=%d "
#                     "(lexical fallback after strict threshold)",
#                     min_hs,
#                     len(reranked),
#                     before - len(reranked),
#                 )
#             else:
#                 logger.info(
#                     "retrieve_rag_pipeline phase=score_threshold min=%.3f kept=0 dropped=%d "
#                     "(strict threshold applied; no fallback context)",
#                     min_hs,
#                     before,
#                 )
#                 reranked = []
#     elif min_hs <= 0.0 and reranked:
#         logger.info(
#             "retrieve_rag_pipeline phase=score_threshold min=%.3f skipped (no hybrid score filter)",
#             min_hs,
#         )

#     _log_retrieval_chunk_stage("filtered_chunks", list(reranked))
#     logger.info(
#         "retrieve_rag_pipeline phase=post_filter duration_sec=%.4f rows_after_filter=%d",
#         time.perf_counter() - t_post_filter,
#         len(reranked),
#     )

#     t_neighbor = time.perf_counter()
#     nb = int(top_doc_chunk_neighbors_before) if top_doc_chunk_neighbors_before is not None else max(0, int(top_doc_chunk_neighbor_radius))
#     na = int(top_doc_chunk_neighbors_after) if top_doc_chunk_neighbors_after is not None else max(0, int(top_doc_chunk_neighbor_radius))
#     if reranked:
#         reranked = expand_reranked_with_doc_chunk_neighbors(
#             reranked,
#             metadata,
#             neighbor_before=nb,
#             neighbor_after=na,
#         )

#     _log_retrieval_chunk_stage("remaining_retrieved_chunks", list(reranked))
#     logger.info(
#         "retrieve_rag_pipeline phase=neighbor_expand duration_sec=%.4f rows_after_expand=%d neighbors_before=%d neighbors_after=%d",
#         time.perf_counter() - t_neighbor,
#         len(reranked),
#         nb,
#         na,
#     )

#     t_pack = time.perf_counter()
#     retrieved = [ch for ch, _ in reranked]
#     if restrict_top_document:
#         capped = pack_top_doc_neighbor_pages_context(
#             reranked,
#             metadata,
#             max_chars=max_context_chars,
#             page_window=top_doc_page_window,
#             max_chunks=max_chunks_for_prompt if max_chunks_for_prompt > 0 else None,
#             question=question,
#         )
#     else:
#         capped = pack_context_from_reranked(
#             reranked,
#             max_context_chars,
#             max_chunks=max_chunks_for_prompt if max_chunks_for_prompt > 0 else None,
#         )
#     capped = ensure_non_empty_capped_context(capped, retrieved, max_context_chars)
#     ctx_chars = sum(len(str(c.get("text", ""))) for c in capped)
#     top_score = float(reranked[0][1]) if reranked else 0.0

#     # Log a compact view of the final context passed to the LLM.
#     if capped:
#         top_doc = str(capped[0].get("doc_name", "") or "")
#         pages = sorted({p for p in (_parse_page_number(c.get("page_number")) for c in capped) if p is not None})
#         logger.info(
#             "final_llm_context doc=%s chunks=%d pages=%s context_chars=%d",
#             top_doc,
#             len(capped),
#             ",".join(str(p) for p in pages) if pages else "?",
#             ctx_chars,
#         )
#         for i, ch in enumerate(capped[: min(8, len(capped))], start=1):
#             logger.info("final_llm_context_chunk #%d %s", i, _format_chunk_for_retrieval_log(ch))
#         if len(capped) > 8:
#             logger.info("final_llm_context_chunk ... (%d more chunks omitted)", len(capped) - 8)
#     logger.info(
#         "retrieve_rag_pipeline phase=context_pack duration_sec=%.4f capped_chunks=%d context_chars=%d",
#         time.perf_counter() - t_pack,
#         len(capped),
#         ctx_chars,
#     )
#     logger.info(
#         "retrieve_rag_pipeline phase=done duration_sec=%.4f corpus_chunks=%d pool_k=%d initial_k=%d top_k=%d "
#         "bm25_weight=%.2f rrf_blend=%.2f retrieved=%d capped_chunks=%d max_llm_chunks=%d context_chars=%d "
#         "top_hybrid_score=%.4f restrict_top_doc=%s backend=%s query_preview=%r",
#         time.perf_counter() - t_pipe,
#         n,
#         pool_k,
#         initial_k,
#         top_k,
#         w_bm25,
#         blend,
#         len(retrieved),
#         len(capped),
#         max_chunks_for_prompt,
#         ctx_chars,
#         top_score,
#         str(restrict_top_document),
#         backend,
#         pipeline_log_preview(question, max_chars=500),
#     )
#     return retrieved, capped, reranked


# _STOPWORDS = {
#     "a",
#     "an",
#     "the",
#     "is",
#     "are",
#     "was",
#     "were",
#     "to",
#     "of",
#     "in",
#     "on",
#     "for",
#     "and",
#     "or",
#     "with",
#     "by",
#     "what",
#     "who",
#     "which",
#     "when",
#     "where",
#     "why",
#     "how",
#     "does",
#     "do",
#     "did",
#     "be",
#     "as",
#     "at",
#     "from",
#     "that",
#     "this",
#     "it",
#     "into",
# }

# _DOC_KEYWORD_HINTS = [
#     (
#         "Calibration-and-Maintenance-of-Store-Equipment",
#         {"calibration", "calibrate", "maintenance", "equipment", "pms", "preventive", "breakdown"},
#     ),
#     (
#         "AccessControlandSecurityofStoreArea",
#         {"access", "security", "visitor", "cctv", "entry", "surveillance", "id card"},
#     ),
#     (
#         "Change-Control-and-Deviation-Management-in-Stores",
#         {"change control", "deviation", "capa", "impact analysis"},
#     ),
#     (
#         "Handling-and-Storage-of-Controlled-Substances",
#         {"controlled substances", "ndps", "narcotic", "drug storage"},
#     ),
#     (
#         "Handling-of-Expired-and-Obsolete-Materials",
#         {"expired", "obsolete", "quarantine", "disposal"},
#     ),
# ]

# def _normalize_token(tok: str) -> str:
#     """
#     Lightweight normalization to improve lexical overlap:
#     - lowercase alnum token
#     - collapse common morphology (calibrate/calibration/calibrated -> calibrat)
#     """
#     t = (tok or "").lower().strip()
#     if len(t) < 2:
#         return ""
#     # Keep short technical tokens unchanged.
#     if len(t) <= 4:
#         return t
#     # Simple suffix stripping (conservative).
#     for suf in (
#         "ization",
#         "isation",
#         "ational",
#         "ating",
#         "ation",
#         "ments",
#         "ment",
#         "ingly",
#         "ness",
#         "able",
#         "ible",
#         "ized",
#         "ised",
#         "izer",
#         "iser",
#         "ical",
#         "ally",
#         "ing",
#         "ers",
#         "ies",
#         "ied",
#         "ed",
#         "es",
#         "s",
#     ):
#         if t.endswith(suf) and len(t) - len(suf) >= 4:
#             t = t[: -len(suf)]
#             break
#     return t


# def _tokenize_for_overlap(text: str) -> List[str]:
#     toks = re.findall(r"[a-zA-Z0-9]+", text.lower())
#     out: List[str] = []
#     for tok in toks:
#         if tok in _STOPWORDS:
#             continue
#         norm = _normalize_token(tok)
#         if not norm or norm in _STOPWORDS or len(norm) < 2:
#             continue
#         out.append(norm)
#     return out


# # def _chunk_body_content_match_score(question: str, seed_blob: str, chunk_text: str) -> float:
# #     """
# #     Lexical overlap in ~[0, 1]: how well chunk body matches the query and/or retrieved seed text.
# #     Used to pull same-doc passages with no section_path_str into context when section expansion skips them.
# #     """
# #     qt = set(_tokenize_for_overlap(question))
# #     st = set(_tokenize_for_overlap(seed_blob))
# #     ct = set(_tokenize_for_overlap(chunk_text))
# #     if not ct:
# #         return 0.0
# #     parts: List[float] = []
# #     if qt:
# #         parts.append(len(qt & ct) / max(1, len(qt)))
# #     if st:
# #         parts.append(len(st & ct) / max(1, len(st)))
# #     if not parts:
# #         return 0.0
# #     return float(sum(parts) / len(parts))


# def hybrid_rerank(
#     question: str,
#     candidate_chunks: List[Dict[str, Any]],
#     candidate_scores: np.ndarray,
#     top_k: int,
#     alpha_semantic: float = 0.72,
#     beta_lexical: float = 0.28,
# ) -> List[Tuple[Dict[str, Any], float]]:
#     """
#     Combine normalized semantic pool scores with token overlap between the question
#     and each chunk (doc name, section path, body). No intent- or domain-specific bonuses.
#     """
#     if not candidate_chunks:
#         return []

#     q_tokens = set(_tokenize_for_overlap(question))
#     if not q_tokens:
#         pairs = [(ch, float(s)) for ch, s in zip(candidate_chunks, candidate_scores)]
#         pairs.sort(key=lambda x: x[1], reverse=True)
#         return pairs[:top_k]

#     sem = np.asarray(candidate_scores, dtype=np.float32)
#     sem_min = float(sem.min())
#     sem_max = float(sem.max())
#     if sem_max > sem_min:
#         sem_norm = (sem - sem_min) / (sem_max - sem_min)
#     else:
#         sem_norm = np.ones_like(sem, dtype=np.float32)

#     lex_scores: List[float] = []
#     for ch in candidate_chunks:
#         text = str(ch.get("text", ""))
#         doc = str(ch.get("doc_name", ""))
#         section = str(ch.get("section_path_str", ""))
#         lexical_text = f"{doc} {section} {text}".strip()
#         c_tokens = set(_tokenize_for_overlap(lexical_text))
#         if not c_tokens:
#             lex_scores.append(0.0)
#             continue
#         overlap = len(q_tokens & c_tokens)
#         lex_scores.append(overlap / max(1, len(q_tokens)))

#     lex = np.asarray(lex_scores, dtype=np.float32)
#     combined = alpha_semantic * sem_norm + beta_lexical * lex
#     ql = (question or "").lower()
#     wants_exact_spec = any(k in ql for k in ("dimension", "range", "rating", "specification", "supply", "capacity", "width", "length", "height"))
#     if wants_exact_spec:
#         for i, ch in enumerate(candidate_chunks):
#             text = str(ch.get("text", "") or "")
#             section = str(ch.get("section_path_str", "") or "")
#             blob = f"{section}\n{text}".lower()
#             bonus = 0.0
#             if "|" in text and re.search(r"\d", text):
#                 bonus += 0.22
#             if any(k in blob for k in ("technical parameter", "technical specification", "main technical", "parameters")):
#                 bonus += 0.18
#             if "carton" in ql and "dimension" in ql and "carton dimension" in blob:
#                 bonus += 0.35
#             if re.search(r"\(\s*\d+\s*[-–]\s*\d+\s*\)\s*mm", text, re.I):
#                 bonus += 0.15
#             if bonus:
#                 combined[i] += bonus
#     order = np.argsort(-combined)
#     out: List[Tuple[Dict[str, Any], float]] = []
#     for i in order[:top_k]:
#         out.append((candidate_chunks[int(i)], float(combined[int(i)])))
#     return out

# def infer_preferred_doc_from_query(question: str, candidate_chunks: List[Dict[str, Any]]) -> Optional[Tuple[str, int]]:
#     """
#     Infer a preferred SOP/document based on query keywords.
#     Returns (doc_name, keyword_hit_strength) when a candidate doc matches a hint.
#     """
#     q = question.lower()
#     if not candidate_chunks:
#         return None

#     candidate_doc_names = {str(ch.get("doc_name", "")) for ch in candidate_chunks}
#     best_doc: Optional[str] = None
#     best_score = 0
#     for doc_hint, kws in _DOC_KEYWORD_HINTS:
#         score = 0
#         for kw in kws:
#             if kw in q:
#                 score += 1
#         if score > best_score:
#             # Find a candidate doc containing the hint token.
#             matched = [d for d in candidate_doc_names if doc_hint.lower() in d.lower()]
#             if matched:
#                 best_doc = matched[0]
#                 best_score = score
#     if not best_doc or best_score <= 0:
#         return None
#     return (best_doc, best_score)

# def enforce_single_doc_scope(
#     question: str,
#     reranked: List[Tuple[Dict[str, Any], float]],
#     top_k: int,
# ) -> List[Tuple[Dict[str, Any], float]]:
#     """
#     Prevent aggressive cross-document filtering.

#     Behavior:
#     1) If query hints a preferred document, keep that document.
#     2) Otherwise, keep mixed top-k results (no forced dominant-doc collapse).

#     This avoids dropping the truly relevant chunks when one unrelated document
#     has slightly higher semantic scores.
#     """
#     if not reranked:
#         return reranked

#     candidates = [ch for ch, _ in reranked]
#     hint = infer_preferred_doc_from_query(question, candidates)
#     if hint:
#         preferred, strength = hint
#         # Single-keyword hints are noisy; avoid wiping a strong mixed retrieval.
#         if strength < 2:
#             return reranked[:top_k]
#         top_doc = str(reranked[0][0].get("doc_name", ""))
#         scoped = [(ch, s) for ch, s in reranked if str(ch.get("doc_name", "")) == preferred]
#         # Narrow only when the best reranked chunk already agrees with the keyword hint.
#         if scoped and top_doc == preferred:
#             return scoped[:top_k]

#     return reranked[:top_k]


# # (Validation / fallback helpers removed. The app now always returns the LLM output.)
import json
import logging
import os
import re
import time
from collections import Counter
from typing import Any, Dict, FrozenSet, Iterable, List, Optional, Set, Tuple

from store_runtime_config import merge_store_runtime_config

# Use the same logger family as the rest of the app so retrieval-chunk logs
# end up in `data/logs/YYYY-MM-DD.log` (configured in qna_service.py).
logger = logging.getLogger("yuktra_qna.rag")


def pipeline_log_preview(text: Optional[str], *, max_chars: int = 4000) -> str:
    """Truncate user/query text for safe logging (single-line friendly)."""
    if text is None:
        return ""
    t = str(text).replace("\r\n", "\n").strip()
    if len(t) <= max_chars:
        return t
    return t[: max_chars - 24] + "\n... [log truncated] ..."


def configure_rag_file_logging(log_dir: str = "data/logs", filename: str = "rag_timing.log") -> None:
    """
    Configure project logging. This now writes to a daily file named `YYYY-MM-DD.log`
    inside `log_dir` and appends (no truncation).

    Notes:
    - Each new day automatically switches to a new file even if the app stays running.
    - Kept for backwards compatibility; prefer using ``logger.get_logger``.
    """
    try:
        from logger import get_logger

        # Configure main project logger (and avoid duplicate handlers across Streamlit reruns).
        get_logger("yuktra_qna", level=logging.INFO, log_dir=log_dir, also_console=False)
    except Exception:
        # Very defensive fallback: don't crash the app just because logging couldn't be configured.
        root = logging.getLogger("yuktra_qna")
        if getattr(root, "_rag_file_logging_configured", False):
            return
        os.makedirs(log_dir, exist_ok=True)
        path = os.path.join(log_dir, filename)
        fh = logging.FileHandler(path, mode="a", encoding="utf-8")
        fh.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(name)s %(message)s"))
        root.setLevel(logging.INFO)
        root.addHandler(fh)
        root._rag_file_logging_configured = True  # type: ignore[attr-defined]


def log_llm_generation_duration(
    duration_sec: float,
    *,
    prompt_chars: int = 0,
    answer_chars: int = 0,
    used_fallback: bool = False,
) -> None:
    """Call from the app after llama.cpp / HF generation finishes."""
    logger.info(
        "llm_generation duration_sec=%.4f prompt_chars=%d answer_chars=%d used_extractive_fallback=%s",
        duration_sec,
        prompt_chars,
        answer_chars,
        str(used_fallback),
    )

import numpy as np
from pypdf import PdfReader
from tqdm import tqdm
import faiss  # type: ignore
#
SUPPORTED_TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".json", ".log"}

# EmbeddingGemma (and similar) use asymmetric prompts for retrieval; see HF model card.
EMBEDDING_PROMPT_STYLE_EMBEDDINGGEMMA = "embeddinggemma"
EMBEDDING_PROMPT_STYLE_PLAIN = "plain"


def resolve_embedding_prompt_style(config: Optional[Dict[str, Any]], embedding_model_id: str) -> str:
    """
    `embeddinggemma` enables official query/document prefixes at embed time.
    Legacy vector stores omit this key — keep `plain` for queries so they match indexed vectors.
    """
    key = (config or {}).get("embedding_prompt_style")
    if key in (EMBEDDING_PROMPT_STYLE_EMBEDDINGGEMMA, EMBEDDING_PROMPT_STYLE_PLAIN):
        return str(key)
    return EMBEDDING_PROMPT_STYLE_PLAIN


def apply_embedding_prompt(text: str, *, role: str, style: str) -> str:
    if style != EMBEDDING_PROMPT_STYLE_EMBEDDINGGEMMA:
        return text
    t = (text or "").strip()
    if not t:
        return text
    if role == "query":
        return f"task: search result | query: {t}"
    if role == "document":
        return f"title: none | text: {t}"
    return text


def is_embeddinggemma_model(model_id_or_path: str) -> bool:
    return "embeddinggemma" in (model_id_or_path or "").lower()

def read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts: List[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        parts.append(page_text)
    return "\n".join(parts).strip()


def extract_pdf_pages(path: str) -> List[Tuple[int, str]]:
    """
    Return per-page extracted text for PDFs so ingestion can preserve page numbers.
    Page numbers are 1-indexed.
    """
    reader = PdfReader(path)
    out: List[Tuple[int, str]] = []
    for i, page in enumerate(reader.pages):
        page_text = (page.extract_text() or "").strip()
        if not page_text:
            continue
        out.append((i + 1, page_text))
    return out

def extract_text_from_docx(path: str) -> str:
    # python-docx is an optional dependency at runtime for docx files.
    from docx import Document  # type: ignore

    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    return "\n".join(parts).strip()

def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext == ".docx":
        return extract_text_from_docx(path)
    if ext in SUPPORTED_TEXT_EXTS:
        return read_text_file(path)
    raise ValueError(f"Unsupported document type: {ext} ({path})")

def chunk_text(text: str, chunk_size_chars: int, overlap_chars: int) -> List[str]:
    """
    Character-based chunking (offline-friendly, no tokenizers required).
    For most RAG setups, this is "good enough" when chunk sizes are tuned.
    """
    if chunk_size_chars <= 0:
        raise ValueError("chunk_size_chars must be > 0")
    if overlap_chars < 0:
        raise ValueError("overlap_chars must be >= 0")
    if overlap_chars >= chunk_size_chars:
        raise ValueError("overlap_chars must be < chunk_size_chars")

    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return []

    # Prefer splitting on paragraph boundaries, then fall back to hard slicing.
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []
    current = ""

    def flush():
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for para in paragraphs:
        if len(current) + len(para) + 2 <= chunk_size_chars:
            current = (current + "\n\n" + para).strip() if current else para
            continue

        # Current chunk is full; flush and start new chunk.
        flush()
        if len(para) <= chunk_size_chars:
            current = para
        else:
            # Paragraph itself is too large; slice with overlap.
            start = 0
            while start < len(para):
                end = min(len(para), start + chunk_size_chars)
                piece = para[start:end]
                chunks.append(piece.strip())
                if end == len(para):
                    break
                start = max(0, end - overlap_chars)

    flush()

    # Second pass: enforce overlap between consecutive chunks (best-effort).
    if overlap_chars == 0 or len(chunks) <= 1:
        return chunks
    out: List[str] = [chunks[0]]
    for prev, cur in zip(chunks, chunks[1:]):
        # If overlap is needed, we can trim the start of the new chunk to simulate overlap.
        trimmed = cur
        if len(trimmed) > overlap_chars:
            trimmed = trimmed  # leave as-is; the first pass already respects boundaries
        out.append(trimmed)
    return out

def l2_normalize(x: np.ndarray, axis: int = -1, eps: float = 1e-12) -> np.ndarray:
    norm = np.linalg.norm(x, axis=axis, keepdims=True)
    return x / np.maximum(norm, eps)

def load_llamacpp_embedding_model(
    model_path: str,
    *,
    n_ctx: int = 2048,
    n_threads: Optional[int] = None,
    n_batch: int = 256,
    verbose: bool = False,
    n_gpu_layers: int = 0,
) -> Any:
    """
    Load a local GGUF embedding model using llama.cpp (via llama-cpp-python).

    n_gpu_layers: number of model layers to offload to GPU. 0 = CPU only
    (default; safe everywhere). -1 = offload all layers. Requires
    llama-cpp-python built with GPU support; otherwise the value is ignored
    silently by the library.

    Returns:
      A `llama_cpp.Llama` instance configured for embeddings.
    """
    try:
        from llama_cpp import Llama  # type: ignore
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "llama-cpp-python is required for llama.cpp embeddings. "
            "Install it (e.g. `pip install llama-cpp-python`) and try again."
        ) from e

    if not model_path or not os.path.isfile(model_path):
        raise FileNotFoundError(f"Embedding GGUF not found: {model_path}")

    logger.info(
        "load_llamacpp_embedding_model start path=%s n_ctx=%s n_threads=%s n_batch=%s n_gpu_layers=%s",
        model_path,
        n_ctx,
        n_threads,
        n_batch,
        n_gpu_layers,
    )
    t0 = time.perf_counter()
    llm = Llama(
        model_path=model_path,
        embedding=True,
        n_ctx=int(n_ctx),
        n_threads=int(n_threads) if n_threads else None,
        n_batch=int(n_batch),
        n_gpu_layers=int(n_gpu_layers),
        verbose=bool(verbose),
    )
    logger.info(
        "load_llamacpp_embedding_model done duration_sec=%.4f path=%s",
        time.perf_counter() - t0,
        model_path,
    )
    return llm


def _extract_llamacpp_embeddings(resp: Any) -> np.ndarray:
    """
    llama-cpp-python embedding response normalizer.
    Expected shape:
      {"data": [{"embedding": [..]}, ...]}
    """
    if not isinstance(resp, dict):
        raise RuntimeError("Unexpected llama.cpp embedding response type.")
    data = resp.get("data")
    if not isinstance(data, list) or not data:
        raise RuntimeError("llama.cpp embedding response missing `data`.")
    vecs: List[np.ndarray] = []
    for row in data:
        emb = (row or {}).get("embedding")
        if not isinstance(emb, list) or not emb:
            raise RuntimeError("llama.cpp embedding response row missing `embedding`.")
        vecs.append(np.asarray(emb, dtype=np.float32))
    out = np.vstack([v.reshape(1, -1) for v in vecs]).astype(np.float32)
    return out


def embed_texts_llamacpp(
    texts: List[str],
    llm: Any,
    *,
    batch_size: int = 16,
    embedding_prompt_style: str = EMBEDDING_PROMPT_STYLE_PLAIN,
    embedding_prompt_role: str = "document",
) -> np.ndarray:
    t_embed = time.perf_counter()
    total_chars = sum(len(t or "") for t in texts)
    if not texts:
        return np.zeros((0, 0), dtype=np.float32)

    logger.info(
        "embed_texts_llamacpp start num_texts=%d total_input_chars=%d batch_size=%d role=%s style=%s",
        len(texts),
        total_chars,
        batch_size,
        embedding_prompt_role,
        embedding_prompt_style,
    )
    all_vecs: List[np.ndarray] = []
    for i in tqdm(range(0, len(texts), batch_size), desc="Embedding chunks"):
        batch_raw = texts[i : i + batch_size]
        batch = [
            apply_embedding_prompt(t, role=embedding_prompt_role, style=embedding_prompt_style)
            for t in batch_raw
        ]
        if not hasattr(llm, "create_embedding"):
            raise RuntimeError("llama.cpp embedding model missing `create_embedding` method.")

        # Some llama.cpp embedding models work reliably only with single-input calls.
        # Use per-text calls for stability.
        batch_vecs: List[np.ndarray] = []
        for t in batch:
            resp = llm.create_embedding(input=t)  # type: ignore[attr-defined]
            v = _extract_llamacpp_embeddings(resp)
            if v.shape[0] != 1:
                raise RuntimeError("Expected a single embedding vector per input string.")
            batch_vecs.append(v)
        vecs = np.vstack(batch_vecs).astype(np.float32)
        vecs = l2_normalize(vecs, axis=1)
        all_vecs.append(vecs)

    out = np.vstack(all_vecs).astype(np.float32)
    logger.info(
        "embed_texts_llamacpp duration_sec=%.4f num_texts=%d total_input_chars=%d prompt_role=%s style=%s",
        time.perf_counter() - t_embed,
        len(texts),
        total_chars,
        embedding_prompt_role,
        embedding_prompt_style,
    )
    return out


def save_vector_store(out_dir: str, vectors: np.ndarray, metadata: List[Dict[str, Any]], config: Dict[str, Any]) -> None:
    logger.info(
        "save_vector_store start out_dir=%s num_vectors=%d embedding_dim=%s index_name=%s tenant=%s",
        out_dir,
        int(vectors.shape[0]) if vectors.ndim == 2 else -1,
        int(vectors.shape[1]) if vectors.ndim == 2 else -1,
        config.get("index_name", ""),
        config.get("tenant_name", ""),
    )
    os.makedirs(out_dir, exist_ok=True)
    meta_path = os.path.join(out_dir, "metadata.json")
    cfg_path = os.path.join(out_dir, "config.json")
    faiss_path = os.path.join(out_dir, "index.faiss")

    if vectors.ndim != 2:
        raise ValueError("vectors must be 2D for FAISS export")
    if vectors.dtype != np.float32:
        vectors = vectors.astype(np.float32)

    # Persist FAISS index as the primary vector store artifact.
    dim = int(vectors.shape[1])
    hnsw_m = 32
    index = faiss.IndexHNSWFlat(dim, hnsw_m, faiss.METRIC_INNER_PRODUCT)
    index.hnsw.efConstruction = 200
    index.hnsw.efSearch = 64
    index.add(vectors)

    faiss.write_index(index, faiss_path)
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False)
    cfg_with_faiss = dict(config)
    cfg_with_faiss.update(
        {
            "vector_store_format": "faiss",
            "faiss_index_file": os.path.basename(faiss_path),
            "faiss_index_type": "IndexHNSWFlat",
            "faiss_metric": "inner_product",
            "faiss_hnsw_m": hnsw_m,
            "faiss_ef_construction": 200,
            "faiss_ef_search": 64,
        }
    )
    with open(cfg_path, "w", encoding="utf-8") as f:
        json.dump(cfg_with_faiss, f, ensure_ascii=False, indent=2)
    logger.info(
        "save_vector_store done out_dir=%s wrote_faiss=%s metadata_rows=%d",
        out_dir,
        os.path.basename(faiss_path),
        len(metadata),
    )

def load_vector_store(
    store_dir: str,
) -> Tuple[Optional[np.ndarray], Optional[Any], List[Dict[str, Any]], Dict[str, Any]]:
    """
    Load a vector store directory.

    Returns:
      (vectors, faiss_index, metadata, config)
      Exactly one of vectors or faiss_index is non-None (FAISS-only export uses faiss_index).
    """
    meta_path = os.path.join(store_dir, "metadata.json")
    cfg_path = os.path.join(store_dir, "config.json")
    if not os.path.isdir(store_dir):
        raise FileNotFoundError(f"Vector store directory not found: {store_dir}")
    if not os.path.isfile(meta_path) or not os.path.isfile(cfg_path):
        raise FileNotFoundError(
            f"Vector store incomplete (need metadata.json + config.json): {store_dir}"
        )

    logger.info("load_vector_store start store_dir=%s", store_dir)

    with open(meta_path, "r", encoding="utf-8") as f:
        metadata = json.load(f)
    with open(cfg_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    config = merge_store_runtime_config(config)
    config = attach_corpus_retrieval_vocab(config, metadata)

    vec_path = os.path.join(store_dir, "vectors.npy")
    faiss_name = str(config.get("faiss_index_file") or "index.faiss")
    faiss_path = os.path.join(store_dir, faiss_name)

    if os.path.isfile(faiss_path):
        index = faiss.read_index(faiss_path)
        if hasattr(index, "hnsw") and config.get("faiss_ef_search") is not None:
            try:
                index.hnsw.efSearch = int(config["faiss_ef_search"])
            except Exception:
                pass
        n = int(index.ntotal)
        if len(metadata) != n:
            raise ValueError(
                f"metadata length ({len(metadata)}) != FAISS index size ({n})."
            )
        logger.info(
            "load_vector_store done backend=faiss ntotal=%d dim=%d format=%s",
            n,
            int(getattr(index, "d", 0) or 0),
            config.get("vector_store_format", "faiss"),
        )
        return None, index, metadata, config

    if os.path.isfile(vec_path):
        vectors = np.load(vec_path).astype(np.float32)
        if len(metadata) != vectors.shape[0]:
            raise ValueError(
                f"metadata length ({len(metadata)}) != vectors count ({vectors.shape[0]})."
            )
        logger.info(
            "load_vector_store done backend=numpy rows=%d dim=%d",
            int(vectors.shape[0]),
            int(vectors.shape[1]),
        )
        return vectors, None, metadata, config

    raise FileNotFoundError(
        f"No vector data found in {store_dir} (expected index.faiss or vectors.npy)."
    )


def topk_search_faiss(index: Any, query_vec: np.ndarray, top_k: int) -> Tuple[np.ndarray, np.ndarray]:
    """Inner-product search (vectors expected L2-normalized). Returns (indices, scores) best-first."""
    q = np.ascontiguousarray(query_vec.astype(np.float32).reshape(1, -1))
    ntotal = int(index.ntotal)
    k = min(int(top_k), max(ntotal, 0))
    if k <= 0:
        return np.zeros((0,), dtype=np.int64), np.zeros((0,), dtype=np.float32)
    scores, ids = index.search(q, k)
    row_ids = ids[0].astype(np.int64)
    row_scores = scores[0].astype(np.float32)
    valid = row_ids >= 0
    return row_ids[valid], row_scores[valid]


def _faiss_reconstruct_rows(index: Any, row_ids: np.ndarray) -> Optional[np.ndarray]:
    """Stack reconstruct(i) for each id; None if not supported."""
    try:
        rows: List[np.ndarray] = []
        for i in np.asarray(row_ids, dtype=np.int64).reshape(-1):
            rows.append(np.asarray(index.reconstruct(int(i)), dtype=np.float32))
        if not rows:
            return np.zeros((0, index.d), dtype=np.float32)
        return np.vstack(rows).astype(np.float32)
    except Exception:
        return None


def topk_search(vectors: np.ndarray, query_vec: np.ndarray, top_k: int) -> Tuple[np.ndarray, np.ndarray]:
    """
    Cosine similarity via dot product (vectors are expected normalized).
    Returns (indices, scores) sorted by descending score.
    """
    if vectors.ndim != 2:
        raise ValueError("vectors must be 2D")
    if query_vec.ndim == 2 and query_vec.shape[0] == 1:
        query_vec = query_vec[0]
    if query_vec.ndim != 1:
        raise ValueError("query_vec must be 1D")

    scores = vectors @ query_vec.astype(np.float32)  # [N]
    if top_k >= len(scores):
        idx = np.argsort(-scores)
    else:
        # partial sort for speed, then fully sort top-k slice
        idx = np.argpartition(-scores, top_k)[:top_k]
        idx = idx[np.argsort(-scores[idx])]
    return idx, scores[idx]


def mmr_select_indices(
    pool_row_indices: np.ndarray,
    vectors: np.ndarray,
    query_vec: np.ndarray,
    k_out: int,
    lambda_mult: float = 0.62,
) -> Tuple[np.ndarray, np.ndarray]:
    """
    Maximal Marginal Relevance on a candidate pool (row indices into `vectors`).
    Returns (ordered_row_indices, scores[dot with query]) sorted by semantic score descending.
    """
    pool_row_indices = np.asarray(pool_row_indices, dtype=np.int64).reshape(-1)
    if pool_row_indices.size == 0:
        return pool_row_indices, np.zeros((0,), dtype=np.float32)
    q = query_vec.astype(np.float32).reshape(-1)
    if pool_row_indices.size <= k_out:
        sel = pool_row_indices
        sc = vectors[sel] @ q
        order = np.argsort(-sc)
        return sel[order], sc[order]

    cand_vecs = vectors[pool_row_indices].astype(np.float32)
    sim_q = cand_vecs @ q
    sim_mx = cand_vecs @ cand_vecs.T
    m = cand_vecs.shape[0]
    k_out = min(k_out, m)

    selected_local: List[int] = []
    selected_set = set()
    first = int(np.argmax(sim_q))
    selected_local.append(first)
    selected_set.add(first)

    while len(selected_local) < k_out:
        best_j = -1
        best_mmr = -1e9
        for j in range(m):
            if j in selected_set:
                continue
            max_sim_sel = max(float(sim_mx[j, i]) for i in selected_local)
            mmr = lambda_mult * float(sim_q[j]) - (1.0 - lambda_mult) * max_sim_sel
            if mmr > best_mmr:
                best_mmr = mmr
                best_j = j
        selected_local.append(best_j)
        selected_set.add(best_j)

    sel_rows = pool_row_indices[np.array(selected_local, dtype=np.int64)]
    sc = vectors[sel_rows] @ q
    order = np.argsort(-sc)
    return sel_rows[order], sc[order]


def pack_context_from_reranked(
    reranked: List[Tuple[Dict[str, Any], float]],
    max_chars: int,
    min_slice_chars: int = 280,
    max_chunks: Optional[int] = None,
) -> List[Dict[str, Any]]:
    """
    Fill the LLM context budget in rerank order, but skip oversized chunks when the
    budget is low so later smaller (still relevant) chunks can be included.
    """
    if max_chars <= 0:
        return []
    if max_chunks is not None and max_chunks > 0:
        reranked = reranked[:max_chunks]
    out: List[Dict[str, Any]] = []
    remaining = max_chars
    for ch, _ in reranked:
        text = str(ch.get("text", ""))
        if not text.strip():
            continue
        if len(text) <= remaining:
            out.append(ch)
            remaining -= len(text)
            if remaining <= 0:
                break
            continue
        if remaining >= min_slice_chars:
            partial = dict(ch)
            partial["text"] = text[:remaining]
            out.append(partial)
            break
    return out


def _parse_page_number(value: Any) -> Optional[int]:
    if isinstance(value, int):
        return value if value > 0 else None
    s = str(value or "").strip()
    if not s:
        return None
    m = re.search(r"\d+", s)
    if not m:
        return None
    try:
        p = int(m.group(0))
        return p if p > 0 else None
    except Exception:
        return None


def _parse_chunk_index(value: Any) -> Optional[int]:
    if value is None:
        return None
    try:
        ci = int(value)
        return ci if ci >= 0 else None
    except (TypeError, ValueError):
        return None


def expand_reranked_with_doc_chunk_neighbors(
    reranked: List[Tuple[Dict[str, Any], float]],
    metadata: List[Dict[str, Any]],
    *,
    neighbor_before: int = 5,
    neighbor_after: int = 1,
) -> List[Tuple[Dict[str, Any], float]]:
    """
    After hybrid rerank and score filtering, widen each remaining hit along the same
    document using chunk_index neighbors in a directional window
    (current-neighbor_before..current+neighbor_after).
    Seed chunks stay first in original order; neighbors append sorted by (doc_name, chunk_index).
    """
    if not reranked or (neighbor_before <= 0 and neighbor_after <= 0):
        return list(reranked)

    by_doc_ci: Dict[Tuple[str, int], Dict[str, Any]] = {}
    for m in metadata:
        doc = str(m.get("doc_name", "") or "")
        ci = _parse_chunk_index(m.get("chunk_index"))
        if not doc or ci is None:
            continue
        by_doc_ci[(doc, ci)] = m

    def _row_key(ch: Dict[str, Any]) -> Tuple[str, int, int]:
        doc = str(ch.get("doc_name", "") or "")
        vec_id = int(ch.get("vector_id", -1) or -1)
        ci = _parse_chunk_index(ch.get("chunk_index"))
        return (doc, vec_id, ci if ci is not None else -1)

    out: List[Tuple[Dict[str, Any], float]] = []
    seen: set[Tuple[str, int, int]] = set()
    for ch, sc in reranked:
        rk = _row_key(ch)
        if rk in seen:
            continue
        seen.add(rk)
        out.append((ch, float(sc)))

    want_by_doc: Dict[str, set[int]] = {}
    for ch, _ in reranked:
        doc = str(ch.get("doc_name", "") or "")
        ci = _parse_chunk_index(ch.get("chunk_index"))
        if not doc or ci is None:
            continue
        s = want_by_doc.setdefault(doc, set())
        for d in range(-max(0, int(neighbor_before)), max(0, int(neighbor_after)) + 1):
            nj = ci + d
            if nj >= 0:
                s.add(nj)

    extra: List[Tuple[Dict[str, Any], float]] = []
    for doc in sorted(want_by_doc.keys()):
        for ci in sorted(want_by_doc[doc]):
            m = by_doc_ci.get((doc, ci))
            if m is None:
                continue
            rk = _row_key(m)
            if rk in seen:
                continue
            seen.add(rk)
            extra.append((m, 0.0))

    if extra:
        logger.info(
            "expand_reranked_with_doc_chunk_neighbors seeds=%d added_neighbors=%d before=%d after=%d",
            len(out),
            len(extra),
            max(0, int(neighbor_before)),
            max(0, int(neighbor_after)),
        )
    return out + extra


def expand_reranked_with_doc_section_neighbors(
    reranked: List[Tuple[Dict[str, Any], float]],
    metadata: List[Dict[str, Any]],
) -> List[Tuple[Dict[str, Any], float]]:
    """
    For each seed in ``reranked``, pull every other chunk from the same document
    that shares the seed's ``section_path_str``. Useful when a top-1 hit lands in
    a multi-chunk section (e.g. "IX. Machine Orientation and Cleaning") and you
    want the full section in front of the LLM rather than only the highest-scored
    chunk inside it.

    Seeds keep their hybrid scores and original order at the front of the list.
    Section-mates append after, ordered by (doc_name, chunk_index, vector_id) with
    score=0.0 (same convention as the chunk-index neighbor expansion).

    Chunks with an empty ``section_path_str`` are ignored (no expansion key).
    """
    if not reranked:
        return list(reranked)

    by_doc_section: Dict[Tuple[str, str], List[Dict[str, Any]]] = {}
    for m in metadata:
        doc = str(m.get("doc_name", "") or "")
        sec = str(m.get("section_path_str", "") or "").strip()
        if not doc or not sec:
            continue
        by_doc_section.setdefault((doc, sec), []).append(m)

    def _row_key(ch: Dict[str, Any]) -> Tuple[str, int, int]:
        doc = str(ch.get("doc_name", "") or "")
        vec_id = int(ch.get("vector_id", -1) or -1)
        ci = _parse_chunk_index(ch.get("chunk_index"))
        return (doc, vec_id, ci if ci is not None else -1)

    out: List[Tuple[Dict[str, Any], float]] = []
    seen: set[Tuple[str, int, int]] = set()
    for ch, sc in reranked:
        rk = _row_key(ch)
        if rk in seen:
            continue
        seen.add(rk)
        out.append((ch, float(sc)))

    seed_sections: set[Tuple[str, str]] = set()
    for ch, _ in reranked:
        doc = str(ch.get("doc_name", "") or "")
        sec = str(ch.get("section_path_str", "") or "").strip()
        if doc and sec:
            seed_sections.add((doc, sec))

    extra: List[Tuple[Dict[str, Any], float]] = []
    for key in sorted(seed_sections):
        section_chunks_sorted = sorted(
            by_doc_section.get(key, []),
            key=lambda m: (
                _parse_chunk_index(m.get("chunk_index")) if _parse_chunk_index(m.get("chunk_index")) is not None else 10**9,
                int(m.get("vector_id", -1) or -1),
            ),
        )
        for m in section_chunks_sorted:
            rk = _row_key(m)
            if rk in seen:
                continue
            seen.add(rk)
            extra.append((m, 0.0))

    if extra:
        logger.info(
            "expand_reranked_with_doc_section_neighbors seeds=%d added_section_mates=%d sections=%d",
            len(out),
            len(extra),
            len(seed_sections),
        )
    return out + extra


def _preview_text_for_log(text: str, *, head: int = 140, tail: int = 90) -> str:
    t = " ".join((text or "").split())
    if not t:
        return ""
    if len(t) <= head + tail + 5:
        return t
    return f"{t[:head]} … {t[-tail:]}"


def _format_chunk_for_retrieval_log(ch: Dict[str, Any]) -> str:
    doc = str(ch.get("doc_name", "") or "")
    vec_id = ch.get("vector_id", "?")
    chunk_idx = ch.get("chunk_index", "?")
    page = ch.get("page_number", "?")
    sec = str(ch.get("section_path_str", "") or "")
    sec = " ".join(sec.split())
    if len(sec) > 80:
        sec = sec[:77] + "…"
    txt = _preview_text_for_log(str(ch.get("text", "") or ""))
    return (
        f"doc={doc} page={page} chunk={chunk_idx} vec={vec_id}"
        + (f" section={sec}" if sec else "")
        + (f" text=\"{txt}\"" if txt else " text=\"\"")
    )


def _log_retrieval_chunk_stage(
    label: str,
    rows: List[Tuple[Dict[str, Any], float]],
    *,
    max_items: int = 256,
) -> None:
    """One line per chunk for a retrieval phase (initial / filtered / remaining, etc.)."""
    n = len(rows)
    logger.info("retrieve_rag_pipeline %s count=%d", label, n)
    cap = min(n, max(0, int(max_items)))
    for i in range(cap):
        ch, sc = rows[i]
        logger.info(
            "retrieve_rag_pipeline %s #%d score=%.4f %s",
            label,
            i + 1,
            float(sc),
            _format_chunk_for_retrieval_log(ch),
        )
    if n > cap:
        logger.info("retrieve_rag_pipeline %s ... %d more omitted (max_items=%d)", label, n - cap, max_items)


def pack_top_doc_neighbor_pages_context(
    reranked: List[Tuple[Dict[str, Any], float]],
    metadata: List[Dict[str, Any]],
    *,
    max_chars: int,
    page_window: int = 1,
    max_chunks: Optional[int] = None,
    question: str = "",
) -> List[Dict[str, Any]]:
    """
    Build context from the top document only: every reranked row for `doc_name` of
    the first hit, in list order (hybrid-ranked seeds followed by same-doc neighbors
    from `expand_reranked_with_doc_chunk_neighbors`), capped to `max_chars` and optional
    `max_chunks`.

    `metadata` and `page_window` are kept for backward compatibility with call sites;
    neighbor expansion is done in `retrieve_rag_pipeline` before this runs.
    """
    if max_chars <= 0 or not reranked:
        return []

    top_doc = str(reranked[0][0].get("doc_name", "") or "")
    if not top_doc:
        return pack_context_from_reranked(reranked, max_chars, max_chunks=max_chunks)

    ordered: List[Dict[str, Any]] = []
    seen_keys: set[Tuple[str, int, int]] = set()

    def _k(ch: Dict[str, Any]) -> Tuple[str, int, int]:
        doc = str(ch.get("doc_name", "") or "")
        chunk_idx = int(ch.get("chunk_index", -1) or -1)
        vec_id = int(ch.get("vector_id", -1) or -1)
        return (doc, vec_id, chunk_idx)

    # Top document only: collect unique chunks first.
    for ch, _ in reranked:
        if str(ch.get("doc_name", "") or "") != top_doc:
            continue
        key = _k(ch)
        if key in seen_keys:
            continue
        seen_keys.add(key)
        ordered.append(ch)

    if not ordered:
        return pack_context_from_reranked(reranked, max_chars, max_chunks=max_chunks)

    # Deterministic chronological/manual flow for the LLM:
    # page asc -> chunk_index asc -> vector_id asc.
    # This avoids seed-first ordering like 19,18,20... and keeps context as 18,19,20,21...
    def _llm_context_order_key(ch: Dict[str, Any]) -> Tuple[int, int, int]:
        p = _parse_page_number(ch.get("page_number"))
        ci = _parse_chunk_index(ch.get("chunk_index"))
        try:
            vid = int(ch.get("vector_id", -1) or -1)
        except Exception:
            vid = -1
        return (
            p if p is not None else 10**9,
            ci if ci is not None else 10**9,
            vid,
        )

    ordered.sort(key=_llm_context_order_key)

    out: List[Dict[str, Any]] = []
    remaining = max_chars
    for ch in ordered:
        text = str(ch.get("text", "") or "")
        if not text.strip():
            continue
        if max_chunks is not None and max_chunks > 0 and len(out) >= max_chunks:
            break
        if len(text) <= remaining:
            out.append(ch)
            remaining -= len(text)
            if remaining <= 0:
                break
            continue
        if remaining >= 280:
            partial = dict(ch)
            partial["text"] = text[:remaining]
            out.append(partial)
        break
    return out


def retrieval_query_variants(
    question: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> List[str]:
    """
    Lightweight query expansion for embedding fusion (no extra LLM call).
    Improves recall for specs, troubleshooting, and safety questions.
    """
    q = (question or "").strip()
    if not q:
        return ["equipment manual"]
    base = build_retrieval_query(q, clean_enabled=clean_enabled, generic_terms=generic_terms)
    if not base:
        base = q
    variants = [base]
    if clean_enabled and base.lower() != q.lower():
        variants.append(q)
    ql = q.lower()

    if any(
        k in ql
        for k in (
            "voltage",
            "volt",
            "vac",
            "vdc",
            "electrical",
            "power supply",
            "line voltage",
            "mains",
            "frequency",
            "hz",
            "amp",
            "current",
        )
    ):
        variants.append(
            f"{base} rated input voltage AC DC electrical specifications power consumption frequency"
        )

    if any(
        k in ql
        for k in (
            "troubleshoot",
            "troubleshooting",
            "not light",
            "does not light",
            "indicator",
            "power indicator",
            "lamp",
            "fault",
            "not working",
            "doesn't work",
            "won't turn",
        )
    ):
        variants.append(
            f"{base} troubleshooting diagnostic fault repair power connection indicator lamp supply"
        )

    if any(
        k in ql
        for k in ("safe", "safety", "hazard", "precaution", "warning", "ppe", "protective", "eye")
    ):
        variants.append(f"{base} safety hazard warning precaution operation protective equipment")

    if any(k in ql for k in ("weight", "kg", "dimension", "size", "mass")):
        variants.append(f"{base} physical specifications weight dimensions")

    if any(
        k in ql
        for k in (
            "calibrat",
            "verification",
            "qualification",
            "validation protocol",
            " as-found",
            "as-left",
            "iq ",
            "oq ",
            "pq ",
        )
    ):
        variants.append(
            f"{base} calibration verification qualification adjustment tolerance specification check procedure"
        )

    if any(
        k in ql
        for k in (
            "temperature",
            "humidity",
            "dew point",
            "relative humidity",
            "rh ",
            "pressure",
            " psi",
            " bar",
            " kpa",
            " mpa",
        )
    ):
        variants.append(f"{base} environmental operating conditions range limits specifications")

    if any(
        k in ql
        for k in (
            "chromatograph",
            "hplc",
            "uhplc",
            "lc-ms",
            "gc-ms",
            "2d-lc",
            "2d lc",
            "two-dimensional",
            "heart-cut",
            "heart cut",
            "heart-cutting",
            "multiple heart",
            "comprehensive 2d",
            "lc x lc",
            "lclc",
        )
    ):
        variants.append(
            f"{base} separation column valve sample loop injection modulation multidimensional chromatography method "
            "principle backward reverse order contamination flush gradient parking deck peak lost occupied loop analysis"
        )

    if any(
        k in ql
        for k in (
            "what is",
            "define",
            "definition",
            "principle",
            "principles",
            "overview",
            "purpose of",
        )
    ):
        variants.append(f"{base} characteristics description notes operation overview")

    seen: set[str] = set()
    out: List[str] = []
    for v in variants:
        k = v.lower()
        if k not in seen:
            seen.add(k)
            out.append(v)
    return out


BM25_LEXICAL_QUERY_MAX_CHARS = 900


def build_bm25_lexical_query(
    question: str,
    max_chars: int = BM25_LEXICAL_QUERY_MAX_CHARS,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> str:
    """
    BM25 over the initial retrieval pool: original question plus the first expansion variant
    (from retrieval_query_variants), capped for tokenizer/BM25 stability.
    """
    variants = retrieval_query_variants(
        question, clean_enabled=clean_enabled, generic_terms=generic_terms
    )
    base = (variants[0] if variants else (question or "")).strip()
    if len(variants) <= 1:
        return base[:max_chars] if max_chars > 0 else base
    extra = variants[1].strip()
    combined = f"{base} {extra}".strip()
    if max_chars <= 0 or len(combined) <= max_chars:
        return combined
    return combined[:max_chars]


def embed_fused_query_for_retrieval(
    question: str,
    tokenizer: Any,
    model: Any,
    device: str,
    max_length: int,
    embedding_prompt_style: str,
    batch_size: int = 8,
    retrieval_query_clean_enabled: bool = True,
    retrieval_generic_terms: Optional[Iterable[str]] = None,
) -> np.ndarray:
    """
    Embed query variants and L2-normalize the mean vector (better recall than a single embed).
    """
    t0 = time.perf_counter()
    q_preview = pipeline_log_preview(question, max_chars=2000)
    retrieval_q = build_retrieval_query(
        question,
        clean_enabled=retrieval_query_clean_enabled,
        generic_terms=retrieval_generic_terms,
    )
    retrieval_preview = pipeline_log_preview(retrieval_q, max_chars=2000)
    logger.info(
        "embed_fused_query_for_retrieval start query_chars=%d query_retrieval_chars=%d "
        "clean_enabled=%s generic_terms=%d style=%s max_length=%d device=%s query=%r query_retrieval=%r",
        len(question or ""),
        len(retrieval_q or ""),
        retrieval_query_clean_enabled,
        len(_coerce_generic_terms(retrieval_generic_terms)),
        embedding_prompt_style,
        max_length,
        device,
        q_preview,
        retrieval_preview,
    )
    variants = retrieval_query_variants(
        question,
        clean_enabled=retrieval_query_clean_enabled,
        generic_terms=retrieval_generic_terms,
    )
    bs = max(1, min(batch_size, len(variants)))
    backend = "llamacpp"
    logger.info(
        "embed_fused_query_for_retrieval variants_ready count=%d batch_size=%d backend=%s",
        len(variants),
        bs,
        backend,
    )
    if not hasattr(model, "create_embedding"):
        raise RuntimeError(
            "Embedding backend must be llama.cpp (expected `create_embedding` on model)."
        )
    vecs = embed_texts_llamacpp(
        variants,
        model,
        batch_size=bs,
        embedding_prompt_style=embedding_prompt_style,
        embedding_prompt_role="query",
    )
    fused = np.mean(vecs.astype(np.float32), axis=0)
    out = l2_normalize(fused.reshape(1, -1), axis=1)[0]
    logger.info(
        "embed_fused_query_for_retrieval done duration_sec=%.4f dim=%d backend=%s",
        time.perf_counter() - t0,
        int(out.shape[0]),
        backend,
    )
    return out


def _minmax_norm_1d(x: np.ndarray, eps: float = 1e-8) -> np.ndarray:
    x = np.asarray(x, dtype=np.float32)
    lo = float(x.min())
    hi = float(x.max())
    if hi - lo < eps:
        return np.ones_like(x, dtype=np.float32)
    return ((x - lo) / (hi - lo + eps)).astype(np.float32)


def _tokenize_bm25(s: str) -> List[str]:
    return re.findall(r"[a-zA-Z0-9]+", (s or "").lower())


def bm25_scores_for_pool(
    query: str,
    doc_texts: List[str],
    k1: float = 1.2,
    b: float = 0.75,
) -> np.ndarray:
    """BM25 over the current candidate pool only (no external index)."""
    N = len(doc_texts)
    if N == 0:
        return np.zeros(0, dtype=np.float32)
    tokenized_docs = [_tokenize_bm25(t) for t in doc_texts]
    doc_lens = np.array([max(1, len(td)) for td in tokenized_docs], dtype=np.float32)
    avgdl = float(np.mean(doc_lens)) + 1e-6

    df: Dict[str, int] = {}
    for td in tokenized_docs:
        for t in set(td):
            df[t] = df.get(t, 0) + 1

    q_terms = _tokenize_bm25(query)
    if not q_terms:
        return np.zeros(N, dtype=np.float32)

    scores = np.zeros(N, dtype=np.float32)
    for i, td in enumerate(tokenized_docs):
        tf: Dict[str, int] = {}
        for t in td:
            tf[t] = tf.get(t, 0) + 1
        dl = doc_lens[i]
        s = 0.0
        for qt in q_terms:
            if qt not in tf:
                continue
            dfi = df.get(qt, 0)
            idf = float(np.log((N - dfi + 0.5) / (dfi + 0.5) + 1.0))
            f = float(tf[qt])
            denom = f + k1 * (1.0 - b + b * (dl / avgdl))
            s += idf * (f * (k1 + 1.0)) / (denom + 1e-8)
        scores[i] = s
    return scores


_BOILERPLATE_NOISE = re.compile(
    r"iso\s*9001|quality\s+management\s+system|about\s+this\s+document|"
    r"desktop\s+computer|minimum\s+requirement|ergonomic|keyboard\s+and\s+mouse|"
    r"limitations?\s+of\s+warranty|registered\s+trademarks",
    re.I,
)
_TECH_IN_CHUNK = re.compile(
    r"voltage|vac|vdc|troubleshoot|indicator|12\s*v|24\s*v|power\s+supply|"
    r"lamp\s+module|fuse|wiring|specification|rated\s+input|green|wall\s+outlet|figure",
    re.I,
)


def boilerplate_noise_multiplier(chunk_text: str) -> float:
    """Downweight generic intro / ISO / PC-requirement chunks that lack substantive technical text."""
    if not chunk_text:
        return 1.0
    head = chunk_text[:12000]
    if _BOILERPLATE_NOISE.search(head) and not _TECH_IN_CHUNK.search(head):
        return 0.32
    return 1.0


_RETRIEVAL_Q_STOP = frozenset(
    {
        "what",
        "the",
        "is",
        "are",
        "was",
        "were",
        "does",
        "did",
        "do",
        "how",
        "when",
        "where",
        "why",
        "which",
        "who",
        "this",
        "that",
        "for",
        "with",
        "from",
        "have",
        "has",
        "should",
        "would",
        "could",
        "can",
        "will",
        "been",
        "being",
        "into",
        "about",
        "your",
        "any",
        "all",
        "and",
        "or",
        "not",
        "but",
        "of",
        "in",
        "on",
        "at",
        "to",
        "an",
        "explain",
        "describe",
        "tell",
        "show",
        "give",
        "please",
        "detail",
        "detailed",
    }
)

# Universal low-signal words (domain-agnostic). Corpus-specific terms are derived at store load.
_RETRIEVAL_BASELINE_GENERIC_TERMS: FrozenSet[str] = frozenset(
    {
        "automatic",
        "machine",
        "manual",
        "equipment",
        "device",
        "system",
        "document",
        "procedure",
        "process",
        "overview",
        "instruction",
        "instructions",
        "operation",
        "operations",
        "guide",
        "section",
        "chapter",
        "content",
    }
)

_METADATA_TITLE_SKIP_TOKENS: FrozenSet[str] = frozenset(
    {
        "pdf",
        "mp4",
        "doc",
        "docx",
        "xlsx",
        "pptx",
        "file",
        "page",
        "copy",
        "draft",
        "final",
        "rev",
        "version",
        "img",
        "image",
        "video",
    }
)


def _metadata_title_tokens(doc_name: str, section: str) -> List[str]:
    """Tokens from a document filename and section heading (title-like surfaces)."""
    base = os.path.splitext(os.path.basename(doc_name or ""))[0]
    base = re.sub(r"[\(\)\[\]_\-]+", " ", base)
    out: List[str] = []
    for part in (base, section or ""):
        for tok in _tokenize_bm25(part):
            if len(tok) < 3 or tok.isdigit() or tok in _METADATA_TITLE_SKIP_TOKENS:
                continue
            out.append(tok)
    return out


def build_corpus_generic_terms(
    metadata: List[Dict[str, Any]],
    config: Optional[Dict[str, Any]] = None,
) -> FrozenSet[str]:
    """
    Derive retrieval noise terms from indexed metadata (doc names, section titles,
  thin-header chunks). Adapts per vector store instead of hardcoding equipment vocabulary.
    """
    if not metadata:
        return _RETRIEVAL_BASELINE_GENERIC_TERMS

    generic: Set[str] = set(_RETRIEVAL_BASELINE_GENERIC_TERMS)
    doc_names = {
        str(m.get("doc_name", "") or "").strip()
        for m in metadata
        if str(m.get("doc_name", "") or "").strip()
    }
    n_docs = len(doc_names)
    if n_docs == 0:
        return frozenset(generic)

    token_doc_names: Dict[str, Set[str]] = {}
    title_token_hits: Counter[str] = Counter()
    body_token_hits: Counter[str] = Counter()
    short_header_token_hits: Counter[str] = Counter()
    long_body_token_hits: Counter[str] = Counter()

    n_chunks = len(metadata)
    for m in metadata:
        doc = str(m.get("doc_name", "") or "")
        section = str(m.get("section_path_str", "") or "")
        text = str(m.get("text", "") or "")
        title_tokens = _metadata_title_tokens(doc, section)
        for tok in set(title_tokens):
            title_token_hits[tok] += 1
            token_doc_names.setdefault(tok, set()).add(doc)

        body_tokens = {t for t in _tokenize_bm25(text) if len(t) >= 3}
        for tok in body_tokens:
            body_token_hits[tok] += 1
            if len(text.strip()) >= 120:
                long_body_token_hits[tok] += 1

        if len(text.strip()) < 100:
            for tok in set(title_tokens):
                short_header_token_hits[tok] += 1

    doc_name_thresh = max(2, int(n_docs * 0.4))
    for tok, docs in token_doc_names.items():
        if tok in _RETRIEVAL_Q_STOP:
            continue
        if len(docs) >= doc_name_thresh:
            generic.add(tok)

    for tok, short_hits in short_header_token_hits.items():
        if short_hits < 2 or tok in _RETRIEVAL_Q_STOP:
            continue
        long_hits = long_body_token_hits.get(tok, 0)
        if long_hits < max(2, short_hits * 0.25):
            if title_token_hits.get(tok, 0) >= 2:
                generic.add(tok)

    for tok, hits in title_token_hits.items():
        if tok in _RETRIEVAL_Q_STOP:
            continue
        if hits / max(1, n_chunks) >= 0.12:
            body_hits = body_token_hits.get(tok, 0)
            if body_hits / max(1, hits) < 0.45:
                generic.add(tok)

    return frozenset(generic)


def attach_corpus_retrieval_vocab(
    config: Dict[str, Any],
    metadata: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """Compute and cache resolved generic terms on the merged store config."""
    if not config.get("retrieval_corpus_vocab_enabled", True):
        resolved = set(_RETRIEVAL_BASELINE_GENERIC_TERMS)
    else:
        resolved = set(build_corpus_generic_terms(metadata, config))
    extras = {str(t).lower().strip() for t in (config.get("retrieval_generic_terms") or []) if str(t).strip()}
    exclude = {str(t).lower().strip() for t in (config.get("retrieval_generic_terms_exclude") or []) if str(t).strip()}
    resolved = (resolved | extras) - exclude
    config["_corpus_generic_terms"] = tuple(sorted(resolved))
    logger.info(
        "attach_corpus_retrieval_vocab corpus_chunks=%d docs=%d generic_terms=%d sample=%r",
        len(metadata),
        len({str(m.get('doc_name', '')) for m in metadata}),
        len(resolved),
        list(sorted(resolved))[:24],
    )
    return config


def resolve_retrieval_generic_terms(config: Optional[Dict[str, Any]] = None) -> FrozenSet[str]:
    cached = (config or {}).get("_corpus_generic_terms")
    if cached is not None:
        return frozenset(str(t).lower() for t in cached)
    return _RETRIEVAL_BASELINE_GENERIC_TERMS


def _coerce_generic_terms(generic_terms: Optional[Iterable[str]] = None) -> FrozenSet[str]:
    if generic_terms is None:
        return _RETRIEVAL_BASELINE_GENERIC_TERMS
    return frozenset(str(t).lower().strip() for t in generic_terms if str(t).strip())

_RETRIEVAL_FOCUS_PATTERNS = (
    re.compile(
        r"^(?:what|which)\s+is\s+(?:the\s+)?(.+?)\s+of\s+(?:the\s+)?.+$",
        re.I,
    ),
    re.compile(
        r"^(?:what|which)\s+are\s+(?:the\s+)?(.+?)\s+of\s+(?:the\s+)?.+$",
        re.I,
    ),
    re.compile(
        r"^(?:tell\s+me|explain|describe)\s+(?:the\s+)?(.+?)\s+of\s+(?:the\s+)?.+$",
        re.I,
    ),
)

_MODEL_SECTION_RE = re.compile(r"^model\b", re.I)
_SPEC_SECTION_RE = re.compile(
    r"technical\s+parameters?|specifications?|specs?\b",
    re.I,
)

_GENERIC_SECTION_RE = re.compile(
    r"how\s+to\s+use|operation\s+manual$|^content$|^overview$|^user'?s?\s+manual$",
    re.I,
)


def _retrieval_lexical_terms(
    question: str,
    generic_terms: Optional[Iterable[str]] = None,
) -> List[str]:
    """Content words for lexical grounding; drops stopwords and corpus-wide generic terms."""
    noise = _coerce_generic_terms(generic_terms)
    return [
        t
        for t in re.findall(r"[a-zA-Z]{3,}", (question or "").lower())
        if t not in _RETRIEVAL_Q_STOP and t not in noise
    ]


def _normalize_section_heading(section: str) -> str:
    s = re.sub(r"^#+\s*", "", (section or "").strip().lower())
    s = re.sub(r"^[ivxlc]+\.\s*", "", s)
    s = re.sub(r"[^a-z0-9\s]", " ", s)
    return re.sub(r"\s+", " ", s).strip()


def _retrieval_topic_phrases(
    question: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> List[str]:
    """Consecutive content-word phrases (bigrams/trigrams) for section-title matching."""
    terms = _retrieval_lexical_terms(
        _retrieval_active_question(question, clean_enabled=clean_enabled, generic_terms=generic_terms),
        generic_terms=generic_terms,
    )
    phrases: List[str] = []
    for n in (3, 2):
        for i in range(len(terms) - n + 1):
            phrases.append(" ".join(terms[i : i + n]))
    return phrases


def _extract_retrieval_focus(question: str) -> Optional[str]:
    """
    Pull the asked attribute from patterns like 'what is X of Y machine'.
    Returns None when the query is already a short direct topic.
    """
    raw = re.sub(r"\s+", " ", (question or "").strip())
    if not raw:
        return None
    for pat in _RETRIEVAL_FOCUS_PATTERNS:
        m = pat.match(raw)
        if not m:
            continue
        focus = re.sub(r"[?.!,;:]+$", "", m.group(1).strip())
        if len(focus) >= 3:
            return focus
    return None


def _retrieval_active_question(
    question: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> str:
    """Lexical retrieval view of the query (never used for LLM generation)."""
    if not clean_enabled:
        return (question or "").strip()
    return build_retrieval_query(
        question,
        clean_enabled=True,
        generic_terms=generic_terms,
    )


def _focus_to_retrieval_query(focus: str) -> str:
    """Keep extracted attribute phrases intact; only strip question stopwords."""
    tokens = _tokenize_bm25(focus)
    filtered = [
        t
        for t in tokens
        if t not in _RETRIEVAL_Q_STOP
        and (len(t) >= 3 or (len(t) >= 2 and any(c.isdigit() for c in t)))
    ]
    if filtered:
        return " ".join(filtered)
    return re.sub(r"\s+", " ", focus.strip().lower())


def build_retrieval_query(
    question: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> str:
    """
    Retrieval-only query: drop stopwords and corpus-generic terms.
    The original question is still used for LLM answer generation.
    """
    noise = _coerce_generic_terms(generic_terms)
    raw = (question or "").strip()
    if not raw or not clean_enabled:
        return raw
    focus = _extract_retrieval_focus(raw)
    if focus:
        return _focus_to_retrieval_query(focus)
    terms = _retrieval_lexical_terms(raw, generic_terms=noise)
    if len(terms) >= 2:
        return " ".join(terms)
    tokens = _tokenize_bm25(raw)
    filtered = [
        t
        for t in tokens
        if t not in _RETRIEVAL_Q_STOP
        and t not in noise
        and (len(t) >= 3 or (len(t) >= 2 and any(c.isdigit() for c in t)))
    ]
    if len(filtered) >= 2:
        return " ".join(filtered)
    if len(filtered) == 1:
        return filtered[0]
    return raw


def filter_retrieval_lexical_query(
    query: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> str:
    """Strip low-signal terms from BM25 query; fall back to original when too few remain."""
    return build_retrieval_query(query, clean_enabled=clean_enabled, generic_terms=generic_terms)


def section_title_match_multiplier(
    chunk_meta: Dict[str, Any],
    question: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> float:
    """Boost chunks whose section heading matches a multi-word topic phrase from the query."""
    section = _normalize_section_heading(str(chunk_meta.get("section_path_str", "") or ""))
    if not section:
        return 1.0

    if _extract_retrieval_focus(question) and _MODEL_SECTION_RE.search(section):
        return 0.62

    for phrase in _retrieval_topic_phrases(
        question, clean_enabled=clean_enabled, generic_terms=generic_terms
    ):
        if len(phrase) < 5:
            continue
        if phrase in section or section in phrase:
            return 1.55

    terms = _retrieval_lexical_terms(
        _retrieval_active_question(
            question, clean_enabled=clean_enabled, generic_terms=generic_terms
        ),
        generic_terms=generic_terms,
    )
    if len(terms) >= 2:
        sec_tokens = set(_tokenize_bm25(section))
        hits = sum(1 for t in set(terms) if t in sec_tokens)
        if hits >= 2:
            ratio = hits / max(1, len(set(terms)))
            return float(1.0 + 0.35 * min(1.0, ratio))
    return 1.0


def generic_section_penalty_multiplier(
    chunk_meta: Dict[str, Any],
    question: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> float:
    """Downrank preface/TOC sections when the query targets a specific topic phrase."""
    section = _normalize_section_heading(str(chunk_meta.get("section_path_str", "") or ""))
    if not section:
        return 1.0

    topic_phrases = [
        p
        for p in _retrieval_topic_phrases(
            question, clean_enabled=clean_enabled, generic_terms=generic_terms
        )
        if len(p.split()) >= 2
    ]
    if not topic_phrases:
        return 1.0

    for phrase in topic_phrases:
        if len(phrase) >= 5 and (phrase in section or section in phrase):
            return 1.0

    if _GENERIC_SECTION_RE.search(section):
        return 0.58
    return 1.0


def _dedupe_fusion_pool(
    idx: np.ndarray,
    fused: np.ndarray,
    metadata: List[Dict[str, Any]],
) -> Tuple[np.ndarray, np.ndarray]:
    """Keep one row per (doc_name, chunk_index), preserving highest fused score (idx pre-sorted)."""
    out_idx: List[int] = []
    out_fused: List[float] = []
    seen: set[Tuple[str, int]] = set()
    for j in range(len(idx)):
        i = int(idx[j])
        ch = metadata[i]
        doc = str(ch.get("doc_name", "") or "")
        ci = _parse_chunk_index(ch.get("chunk_index"))
        key = (doc, ci if ci is not None else i)
        if key in seen:
            continue
        seen.add(key)
        out_idx.append(i)
        out_fused.append(float(fused[j]))
    return np.asarray(out_idx, dtype=idx.dtype), np.asarray(out_fused, dtype=np.float32)


def query_term_coverage_multiplier(
    chunk_text: str,
    question: str,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> float:
    """Boost chunks that contain more of the user's content words (lexical grounding)."""
    if not chunk_text or not (question or "").strip():
        return 1.0
    qts = _retrieval_lexical_terms(
        _retrieval_active_question(
            question, clean_enabled=clean_enabled, generic_terms=generic_terms
        ),
        generic_terms=generic_terms,
    )
    if len(qts) < 2:
        return 1.0
    blob = chunk_text.lower()
    hits = sum(1 for t in set(qts) if t in blob)
    ratio = hits / max(1, len(set(qts)))
    return float(0.78 + 0.50 * min(1.0, ratio))


def thin_chunk_multiplier(chunk_text: str) -> float:
    """Downrank very short chunks (often headers/noise)."""
    L = len((chunk_text or "").strip())
    if L < 60:
        return 0.35
    if L < 120:
        return 0.55
    if L < 180:
        return 0.85
    return 1.0


def title_only_chunk_penalty(chunk_meta: Dict[str, Any], chunk_text: str) -> float:
    """Penalize cover/title chunks that lack substantive body text."""
    section = _normalize_section_heading(str(chunk_meta.get("section_path_str", "") or ""))
    text = (chunk_text or "").strip()
    if len(text) < 100 and _MODEL_SECTION_RE.search(section):
        return 0.45
    if len(text) < 80:
        return 0.55
    return 1.0


def spec_section_boost_multiplier(
    chunk_meta: Dict[str, Any],
    question: str,
    *,
    generic_terms: Optional[Iterable[str]] = None,
) -> float:
    """Boost spec-table sections for attribute-style questions."""
    section = str(chunk_meta.get("section_path_str", "") or "")
    if not section or not _SPEC_SECTION_RE.search(section):
        return 1.0
    active = _retrieval_active_question(question, generic_terms=generic_terms)
    if len(_retrieval_lexical_terms(active, generic_terms=generic_terms)) >= 1:
        return 1.22
    return 1.0


def _ranks_descending_scores(scores: np.ndarray) -> np.ndarray:
    """rank[j] = 0 for best (highest score), 1 for second, ..."""
    scores = np.asarray(scores, dtype=np.float32)
    order = np.argsort(-scores)
    ranks = np.empty(len(scores), dtype=np.float32)
    ranks[order] = np.arange(len(scores), dtype=np.float32)
    return ranks


def reciprocal_rank_fusion_score(dense_scores: np.ndarray, bm25_scores: np.ndarray, k: float = 58.0) -> np.ndarray:
    """
    RRF combines two rankers without fragile score scaling (Cormack et al. style).
    """
    rd = _ranks_descending_scores(dense_scores)
    rb = _ranks_descending_scores(bm25_scores)
    return (1.0 / (k + rd) + 1.0 / (k + rb)).astype(np.float32)


def ensure_non_empty_capped_context(
    capped: List[Dict[str, Any]],
    retrieved: List[Dict[str, Any]],
    max_context_chars: int,
) -> List[Dict[str, Any]]:
    """
    If packing skipped everything (e.g. oversized chunks vs tight budget), still pass the
    top retrieved chunk so the LLM is not sent empty context.
    """
    if capped:
        return capped
    if not retrieved:
        return []
    ch = dict(retrieved[0])
    t = str(ch.get("text", ""))
    if not t.strip():
        return []
    budget = max_context_chars if max_context_chars > 0 else 8001
    limit = min(len(t), budget)
    ch["text"] = t[:limit]
    return [ch]


def retrieve_rag_pipeline(
    question: str,
    query_vec: np.ndarray,
    vectors: Optional[np.ndarray],
    metadata: List[Dict[str, Any]],
    *,
    faiss_index: Optional[Any] = None,
    top_k: int,
    retrieval_pool_k: int,
    mmr_k: int,
    mmr_lambda: float = 0.62,
    max_context_chars: int,
    restrict_top_document: bool = False,
    bm25_weight: float = 0.6,
    initial_pool_multiplier: int = 2,
    max_chunks_for_prompt: int = 3,
    rrf_linear_blend: float = 0.25,
    hybrid_alpha_semantic: float = 0.30,
    hybrid_rerank_enabled: bool = False,
    min_hybrid_rerank_score: float = 0.0,
    top_doc_page_window: int = 2,
    top_doc_chunk_neighbor_radius: int = 6,
    top_doc_chunk_neighbors_before: Optional[int] = None,
    top_doc_chunk_neighbors_after: Optional[int] = None,
    top_doc_section_expand: bool = True,
    retrieval_query_clean_enabled: bool = True,
    retrieval_generic_terms: Optional[Iterable[str]] = None,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Tuple[Dict[str, Any], float]]]:
    """
    Returns:
      - retrieved: chunks after hybrid rerank, score threshold, and same-doc chunk_index
        ±N neighbor expansion (via metadata, default N=6).
      - capped_for_prompt: packed to max_context_chars (reduces noise)
      - reranked: (chunk, score) pairs after threshold and neighbor expansion; seeds keep
        hybrid scores, neighbors use 0.0.

    Dense + BM25: min-max linear blend plus Reciprocal Rank Fusion (RRF) for robustness; then
    boilerplate / coverage / thin-chunk multipliers. BM25 uses the original question plus the
    first retrieval expansion variant (capped at BM25_LEXICAL_QUERY_MAX_CHARS).
    """
    t_pipe = time.perf_counter()
    n = len(metadata)
    if n == 0:
        logger.info(
            "retrieve_rag_pipeline empty_store duration_sec=%.4f query_preview=%r",
            time.perf_counter() - t_pipe,
            pipeline_log_preview(question, max_chars=800),
        )
        return [], [], []

    if (vectors is None) == (faiss_index is None):
        raise ValueError("retrieve_rag_pipeline: set exactly one of vectors or faiss_index")

    pool_k = min(n, max(top_k, retrieval_pool_k))
    mult = max(1, int(initial_pool_multiplier))
    if n > 40:
        mult = max(mult, 3)
    initial_k = min(n, max(pool_k * mult, pool_k + min(32, n // 4)))

    backend = "faiss" if faiss_index is not None else "numpy"
    generic_terms = _coerce_generic_terms(retrieval_generic_terms)
    q_log = pipeline_log_preview(question, max_chars=3500)
    retrieval_q = build_retrieval_query(
        question,
        clean_enabled=retrieval_query_clean_enabled,
        generic_terms=generic_terms,
    )
    retrieval_log = pipeline_log_preview(retrieval_q, max_chars=3500)
    logger.info(
        "retrieve_rag_pipeline phase=init query_chars=%d query_retrieval_chars=%d clean_enabled=%s "
        "generic_terms=%d backend=%s corpus_chunks=%d "
        "top_k=%d retrieval_pool_k=%d mmr_k=%d initial_k=%d pool_k=%d max_context_chars=%d "
        "restrict_top_doc=%s query=%r query_retrieval=%r",
        len(question or ""),
        len(retrieval_q or ""),
        retrieval_query_clean_enabled,
        len(generic_terms),
        backend,
        n,
        top_k,
        retrieval_pool_k,
        mmr_k,
        initial_k,
        pool_k,
        max_context_chars,
        restrict_top_document,
        q_log,
        retrieval_log,
    )

    t_dense = time.perf_counter()
    if faiss_index is not None:
        idx, dense_scores = topk_search_faiss(faiss_index, query_vec, initial_k)
    else:
        assert vectors is not None
        idx, dense_scores = topk_search(vectors, query_vec, initial_k)
    texts = [str(metadata[int(i)].get("text", "")) for i in idx]
    logger.info(
        "retrieve_rag_pipeline phase=dense_search duration_sec=%.4f initial_k=%d returned=%d dense_min=%.4f dense_max=%.4f",
        time.perf_counter() - t_dense,
        initial_k,
        len(idx),
        float(np.min(dense_scores)) if len(dense_scores) else 0.0,
        float(np.max(dense_scores)) if len(dense_scores) else 0.0,
    )

    t_fuse = time.perf_counter()
    bm25_query = filter_retrieval_lexical_query(
        build_bm25_lexical_query(
            question,
            clean_enabled=retrieval_query_clean_enabled,
            generic_terms=generic_terms,
        ),
        clean_enabled=retrieval_query_clean_enabled,
        generic_terms=generic_terms,
    )
    bm25_s = bm25_scores_for_pool(bm25_query, texts)
    dn = _minmax_norm_1d(dense_scores)
    bn = _minmax_norm_1d(bm25_s)
    w_bm25 = float(np.clip(bm25_weight, 0.0, 0.95))
    linear = (1.0 - w_bm25) * dn + w_bm25 * bn
    rrf = reciprocal_rank_fusion_score(dense_scores, bm25_s)
    rrf_n = _minmax_norm_1d(rrf)
    blend = float(np.clip(rrf_linear_blend, 0.0, 0.95))
    fused = (1.0 - blend) * linear + blend * rrf_n

    for j in range(len(idx)):
        ch_meta = metadata[int(idx[j])]
        fused[j] *= boilerplate_noise_multiplier(texts[j])
        fused[j] *= query_term_coverage_multiplier(
            texts[j],
            question,
            clean_enabled=retrieval_query_clean_enabled,
            generic_terms=generic_terms,
        )
        fused[j] *= thin_chunk_multiplier(texts[j])
        fused[j] *= title_only_chunk_penalty(ch_meta, texts[j])
        fused[j] *= section_title_match_multiplier(
            ch_meta,
            question,
            clean_enabled=retrieval_query_clean_enabled,
            generic_terms=generic_terms,
        )
        fused[j] *= generic_section_penalty_multiplier(
            ch_meta,
            question,
            clean_enabled=retrieval_query_clean_enabled,
            generic_terms=generic_terms,
        )
        fused[j] *= spec_section_boost_multiplier(ch_meta, question, generic_terms=generic_terms)

    order = np.argsort(-fused)
    idx = idx[order]
    fused = fused[order]
    idx = idx[:pool_k].copy()
    fused = fused[:pool_k].copy()
    idx, fused = _dedupe_fusion_pool(idx, fused, metadata)
    logger.info(
        "retrieve_rag_pipeline phase=hybrid_fusion duration_sec=%.4f pool_k=%d bm25_weight=%.2f rrf_blend=%.2f bm25_query_chars=%d",
        time.perf_counter() - t_fuse,
        pool_k,
        w_bm25,
        blend,
        len(bm25_query),
    )

    t_mmr = time.perf_counter()
    mmr_take = min(mmr_k, pool_k)
    if mmr_take >= pool_k:
        idx_sel = idx
        scores_sel = fused
    else:
        if faiss_index is not None:
            cand_matrix = _faiss_reconstruct_rows(faiss_index, idx)
            if cand_matrix is None or cand_matrix.shape[0] != len(idx):
                idx_sel = idx[:mmr_take]
                scores_sel = fused[:mmr_take].copy()
            else:
                local = np.arange(len(idx), dtype=np.int64)
                idx_local_sel, _ = mmr_select_indices(
                    local, cand_matrix, query_vec, mmr_take, lambda_mult=mmr_lambda
                )
                idx_sel = idx[idx_local_sel.astype(np.int64)]
                lut = {int(idx[j]): float(fused[j]) for j in range(len(idx))}
                scores_sel = np.array([lut[int(i)] for i in idx_sel], dtype=np.float32)
        else:
            assert vectors is not None
            idx_sel, _ = mmr_select_indices(idx, vectors, query_vec, mmr_take, lambda_mult=mmr_lambda)
            lut = {int(idx[j]): float(fused[j]) for j in range(len(idx))}
            scores_sel = np.array([lut[int(i)] for i in idx_sel], dtype=np.float32)

    logger.info(
        "retrieve_rag_pipeline phase=mmr duration_sec=%.4f mmr_take=%d candidate_rows=%d (pool_before_mmr=%d)",
        time.perf_counter() - t_mmr,
        mmr_take,
        len(idx_sel),
        pool_k,
    )

    candidates = [metadata[int(i)] for i in idx_sel]
    candidate_scores = np.asarray(scores_sel, dtype=np.float32)
    _log_retrieval_chunk_stage(
        "initial_chunks",
        [(c, float(candidate_scores[j])) for j, c in enumerate(candidates)],
    )

    t_rerank = time.perf_counter()
    if hybrid_rerank_enabled:
        alpha_sem = float(np.clip(hybrid_alpha_semantic, 0.05, 0.95))
        beta_lex = 1.0 - alpha_sem
        reranked = hybrid_rerank(
            question,
            candidate_chunks=candidates,
            candidate_scores=candidate_scores,
            top_k=top_k,
            alpha_semantic=alpha_sem,
            beta_lexical=beta_lex,
            clean_enabled=retrieval_query_clean_enabled,
            generic_terms=generic_terms,
        )
        phase_name = "hybrid_rerank"
    else:
        pairs = [(candidates[j], float(candidate_scores[j])) for j in range(len(candidates))]
        pairs.sort(key=lambda x: x[1], reverse=True)
        reranked = pairs[:top_k]
        phase_name = "top_fusion_select"

    if reranked:
        logger.info(
            "retrieve_rag_pipeline phase=%s duration_sec=%.4f rows=%d",
            phase_name,
            time.perf_counter() - t_rerank,
            len(reranked),
        )
        for i, (ch, sc) in enumerate(reranked, start=1):
            logger.info(
                "retrieve_rag_pipeline %s #%d doc=%s chunk_index=%s vec=%s score=%.4f",
                phase_name,
                i,
                str(ch.get("doc_name", "?"))[:120],
                ch.get("chunk_index", "?"),
                ch.get("vector_id", "?"),
                float(sc),
            )
    else:
        logger.info(
            "retrieve_rag_pipeline phase=%s duration_sec=%.4f rows=0 (empty)",
            phase_name,
            time.perf_counter() - t_rerank,
        )
    t_post_filter = time.perf_counter()
    if restrict_top_document:
        reranked = enforce_single_doc_scope(question, reranked, top_k=top_k)

    if restrict_top_document and reranked:
        top_doc = str(reranked[0][0].get("doc_name", "") or "")
        if top_doc:
            reranked = [(ch, s) for ch, s in reranked if str(ch.get("doc_name", "")) == top_doc][:top_k]

    min_hs = float(np.clip(min_hybrid_rerank_score, 0.0, 1.0))
    if min_hs > 0.0 and reranked:
        before = len(reranked)
        dropped_hs = [(ch, s) for ch, s in reranked if float(s) < min_hs]
        filtered_hs = [(ch, s) for ch, s in reranked if float(s) >= min_hs]
        if dropped_hs:
            _log_retrieval_chunk_stage("score_filter_dropped_chunks", dropped_hs)
        if filtered_hs:
            reranked = filtered_hs
            logger.info(
                "retrieve_rag_pipeline phase=score_threshold min=%.3f kept=%d dropped=%d",
                min_hs,
                len(reranked),
                before - len(reranked),
            )
        else:
            logger.info(
                "retrieve_rag_pipeline phase=score_threshold min=%.3f kept=0 dropped=%d "
                "(strict threshold applied; no fallback context)",
                min_hs,
                before,
            )
            reranked = []
    elif min_hs <= 0.0 and reranked:
        logger.info(
            "retrieve_rag_pipeline phase=score_threshold min=%.3f skipped (no hybrid score filter)",
            min_hs,
        )

    _log_retrieval_chunk_stage("filtered_chunks", list(reranked))
    logger.info(
        "retrieve_rag_pipeline phase=post_filter duration_sec=%.4f rows_after_filter=%d",
        time.perf_counter() - t_post_filter,
        len(reranked),
    )

    t_neighbor = time.perf_counter()
    nb = int(top_doc_chunk_neighbors_before) if top_doc_chunk_neighbors_before is not None else max(0, int(top_doc_chunk_neighbor_radius))
    na = int(top_doc_chunk_neighbors_after) if top_doc_chunk_neighbors_after is not None else max(0, int(top_doc_chunk_neighbor_radius))
    if reranked:
        reranked = expand_reranked_with_doc_chunk_neighbors(
            reranked,
            metadata,
            neighbor_before=nb,
            neighbor_after=na,
        )

    _log_retrieval_chunk_stage("remaining_retrieved_chunks", list(reranked))
    logger.info(
        "retrieve_rag_pipeline phase=neighbor_expand duration_sec=%.4f rows_after_expand=%d neighbors_before=%d neighbors_after=%d",
        time.perf_counter() - t_neighbor,
        len(reranked),
        nb,
        na,
    )

    if top_doc_section_expand and reranked:
        t_section = time.perf_counter()
        rows_before_section = len(reranked)
        reranked = expand_reranked_with_doc_section_neighbors(reranked, metadata)
        _log_retrieval_chunk_stage("section_expanded_chunks", list(reranked))
        logger.info(
            "retrieve_rag_pipeline phase=section_expand duration_sec=%.4f rows_before=%d rows_after=%d",
            time.perf_counter() - t_section,
            rows_before_section,
            len(reranked),
        )

    t_pack = time.perf_counter()
    retrieved = [ch for ch, _ in reranked]
    if restrict_top_document:
        capped = pack_top_doc_neighbor_pages_context(
            reranked,
            metadata,
            max_chars=max_context_chars,
            page_window=top_doc_page_window,
            max_chunks=max_chunks_for_prompt if max_chunks_for_prompt > 0 else None,
            question=question,
        )
    else:
        capped = pack_context_from_reranked(
            reranked,
            max_context_chars,
            max_chunks=max_chunks_for_prompt if max_chunks_for_prompt > 0 else None,
        )
    capped = ensure_non_empty_capped_context(capped, retrieved, max_context_chars)
    ctx_chars = sum(len(str(c.get("text", ""))) for c in capped)
    top_score = float(reranked[0][1]) if reranked else 0.0

    # Log a compact view of the final context passed to the LLM.
    if capped:
        top_doc = str(capped[0].get("doc_name", "") or "")
        pages = sorted({p for p in (_parse_page_number(c.get("page_number")) for c in capped) if p is not None})
        logger.info(
            "final_llm_context doc=%s chunks=%d pages=%s context_chars=%d",
            top_doc,
            len(capped),
            ",".join(str(p) for p in pages) if pages else "?",
            ctx_chars,
        )
        for i, ch in enumerate(capped[: min(8, len(capped))], start=1):
            logger.info("final_llm_context_chunk #%d %s", i, _format_chunk_for_retrieval_log(ch))
        if len(capped) > 8:
            logger.info("final_llm_context_chunk ... (%d more chunks omitted)", len(capped) - 8)
    logger.info(
        "retrieve_rag_pipeline phase=context_pack duration_sec=%.4f capped_chunks=%d context_chars=%d",
        time.perf_counter() - t_pack,
        len(capped),
        ctx_chars,
    )
    logger.info(
        "retrieve_rag_pipeline phase=done duration_sec=%.4f corpus_chunks=%d pool_k=%d initial_k=%d top_k=%d "
        "bm25_weight=%.2f rrf_blend=%.2f retrieved=%d capped_chunks=%d max_llm_chunks=%d context_chars=%d "
        "top_hybrid_score=%.4f restrict_top_doc=%s backend=%s query_preview=%r",
        time.perf_counter() - t_pipe,
        n,
        pool_k,
        initial_k,
        top_k,
        w_bm25,
        blend,
        len(retrieved),
        len(capped),
        max_chunks_for_prompt,
        ctx_chars,
        top_score,
        str(restrict_top_document),
        backend,
        pipeline_log_preview(question, max_chars=500),
    )
    return retrieved, capped, reranked


_STOPWORDS = {
    "a",
    "an",
    "the",
    "is",
    "are",
    "was",
    "were",
    "to",
    "of",
    "in",
    "on",
    "for",
    "and",
    "or",
    "with",
    "by",
    "what",
    "who",
    "which",
    "when",
    "where",
    "why",
    "how",
    "does",
    "do",
    "did",
    "be",
    "as",
    "at",
    "from",
    "that",
    "this",
    "it",
    "into",
}

_DOC_KEYWORD_HINTS = [
    (
        "Calibration-and-Maintenance-of-Store-Equipment",
        {"calibration", "calibrate", "maintenance", "equipment", "pms", "preventive", "breakdown"},
    ),
    (
        "AccessControlandSecurityofStoreArea",
        {"access", "security", "visitor", "cctv", "entry", "surveillance", "id card"},
    ),
    (
        "Change-Control-and-Deviation-Management-in-Stores",
        {"change control", "deviation", "capa", "impact analysis"},
    ),
    (
        "Handling-and-Storage-of-Controlled-Substances",
        {"controlled substances", "ndps", "narcotic", "drug storage"},
    ),
    (
        "Handling-of-Expired-and-Obsolete-Materials",
        {"expired", "obsolete", "quarantine", "disposal"},
    ),
]

def _normalize_token(tok: str) -> str:
    """
    Lightweight normalization to improve lexical overlap:
    - lowercase alnum token
    - collapse common morphology (calibrate/calibration/calibrated -> calibrat)
    """
    t = (tok or "").lower().strip()
    if len(t) < 2:
        return ""
    # Keep short technical tokens unchanged.
    if len(t) <= 4:
        return t
    # Simple suffix stripping (conservative).
    for suf in (
        "ization",
        "isation",
        "ational",
        "ating",
        "ation",
        "ments",
        "ment",
        "ingly",
        "ness",
        "able",
        "ible",
        "ized",
        "ised",
        "izer",
        "iser",
        "ical",
        "ally",
        "ing",
        "ers",
        "ies",
        "ied",
        "ed",
        "es",
        "s",
    ):
        if t.endswith(suf) and len(t) - len(suf) >= 4:
            t = t[: -len(suf)]
            break
    return t


def _tokenize_for_overlap(text: str) -> List[str]:
    toks = re.findall(r"[a-zA-Z0-9]+", text.lower())
    out: List[str] = []
    for tok in toks:
        if tok in _STOPWORDS:
            continue
        norm = _normalize_token(tok)
        if not norm or norm in _STOPWORDS or len(norm) < 2:
            continue
        out.append(norm)
    return out


# def _chunk_body_content_match_score(question: str, seed_blob: str, chunk_text: str) -> float:
#     """
#     Lexical overlap in ~[0, 1]: how well chunk body matches the query and/or retrieved seed text.
#     Used to pull same-doc passages with no section_path_str into context when section expansion skips them.
#     """
#     qt = set(_tokenize_for_overlap(question))
#     st = set(_tokenize_for_overlap(seed_blob))
#     ct = set(_tokenize_for_overlap(chunk_text))
#     if not ct:
#         return 0.0
#     parts: List[float] = []
#     if qt:
#         parts.append(len(qt & ct) / max(1, len(qt)))
#     if st:
#         parts.append(len(st & ct) / max(1, len(st)))
#     if not parts:
#         return 0.0
#     return float(sum(parts) / len(parts))


def hybrid_rerank(
    question: str,
    candidate_chunks: List[Dict[str, Any]],
    candidate_scores: np.ndarray,
    top_k: int,
    alpha_semantic: float = 0.72,
    beta_lexical: float = 0.28,
    *,
    clean_enabled: bool = True,
    generic_terms: Optional[Iterable[str]] = None,
) -> List[Tuple[Dict[str, Any], float]]:
    """
    Combine normalized semantic pool scores with token overlap between the retrieval-focused
    query and each chunk (section path, body).
    """
    if not candidate_chunks:
        return []

    active_q = _retrieval_active_question(
        question, clean_enabled=clean_enabled, generic_terms=generic_terms
    )
    q_tokens = set(_retrieval_lexical_terms(active_q, generic_terms=generic_terms))
    if not q_tokens:
        q_tokens = set(_tokenize_for_overlap(active_q))
    if not q_tokens:
        pairs = [(ch, float(s)) for ch, s in zip(candidate_chunks, candidate_scores)]
        pairs.sort(key=lambda x: x[1], reverse=True)
        return pairs[:top_k]

    sem = np.asarray(candidate_scores, dtype=np.float32)
    sem_min = float(sem.min())
    sem_max = float(sem.max())
    if sem_max > sem_min:
        sem_norm = (sem - sem_min) / (sem_max - sem_min)
    else:
        sem_norm = np.ones_like(sem, dtype=np.float32)

    lex_scores: List[float] = []
    for ch in candidate_chunks:
        text = str(ch.get("text", ""))
        section = str(ch.get("section_path_str", ""))
        section_tokens = set(_tokenize_for_overlap(section))
        body_tokens = set(_tokenize_for_overlap(text))
        if not section_tokens and not body_tokens:
            lex_scores.append(0.0)
            continue
        sec_overlap = len(q_tokens & section_tokens) / max(1, len(q_tokens))
        body_overlap = len(q_tokens & body_tokens) / max(1, len(q_tokens))
        lex = 0.35 * sec_overlap + 0.65 * body_overlap
        lex_scores.append(lex)

    lex = np.asarray(lex_scores, dtype=np.float32)
    combined = alpha_semantic * sem_norm + beta_lexical * lex
    for j, ch in enumerate(candidate_chunks):
        text = str(ch.get("text", ""))
        combined[j] *= section_title_match_multiplier(
            ch, question, clean_enabled=clean_enabled, generic_terms=generic_terms
        )
        combined[j] *= generic_section_penalty_multiplier(
            ch, question, clean_enabled=clean_enabled, generic_terms=generic_terms
        )
        combined[j] *= thin_chunk_multiplier(text)
        combined[j] *= title_only_chunk_penalty(ch, text)
        combined[j] *= spec_section_boost_multiplier(ch, question, generic_terms=generic_terms)
    order = np.argsort(-combined)
    out: List[Tuple[Dict[str, Any], float]] = []
    for i in order[:top_k]:
        out.append((candidate_chunks[int(i)], float(combined[int(i)])))
    return out

def infer_preferred_doc_from_query(question: str, candidate_chunks: List[Dict[str, Any]]) -> Optional[Tuple[str, int]]:
    """
    Infer a preferred SOP/document based on query keywords.
    Returns (doc_name, keyword_hit_strength) when a candidate doc matches a hint.
    """
    q = question.lower()
    if not candidate_chunks:
        return None

    candidate_doc_names = {str(ch.get("doc_name", "")) for ch in candidate_chunks}
    best_doc: Optional[str] = None
    best_score = 0
    for doc_hint, kws in _DOC_KEYWORD_HINTS:
        score = 0
        for kw in kws:
            if kw in q:
                score += 1
        if score > best_score:
            # Find a candidate doc containing the hint token.
            matched = [d for d in candidate_doc_names if doc_hint.lower() in d.lower()]
            if matched:
                best_doc = matched[0]
                best_score = score
    if not best_doc or best_score <= 0:
        return None
    return (best_doc, best_score)

def enforce_single_doc_scope(
    question: str,
    reranked: List[Tuple[Dict[str, Any], float]],
    top_k: int,
) -> List[Tuple[Dict[str, Any], float]]:
    """
    Prevent aggressive cross-document filtering.

    Behavior:
    1) If query hints a preferred document, keep that document.
    2) Otherwise, keep mixed top-k results (no forced dominant-doc collapse).

    This avoids dropping the truly relevant chunks when one unrelated document
    has slightly higher semantic scores.
    """
    if not reranked:
        return reranked

    candidates = [ch for ch, _ in reranked]
    hint = infer_preferred_doc_from_query(question, candidates)
    if hint:
        preferred, strength = hint
        # Single-keyword hints are noisy; avoid wiping a strong mixed retrieval.
        if strength < 2:
            return reranked[:top_k]
        top_doc = str(reranked[0][0].get("doc_name", ""))
        scoped = [(ch, s) for ch, s in reranked if str(ch.get("doc_name", "")) == preferred]
        # Narrow only when the best reranked chunk already agrees with the keyword hint.
        if scoped and top_doc == preferred:
            return scoped[:top_k]

    return reranked[:top_k]


# (Validation / fallback helpers removed. The app now always returns the LLM output.)
